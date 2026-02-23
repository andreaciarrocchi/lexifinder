#!/usr/bin/env python3
"""
LexiFinder CLI - Hierarchical Analytic Index Generator

Supports PDF, DOCX, and ODT documents with flexible indexing:
- PDF: index by page or paragraph
- DOCX/ODT: index by paragraph

Creates a two-level hierarchical index:
- Level 1: Main concepts (your keywords)
- Level 2: Related sub-entries (semantically similar nouns)
"""

import argparse
import sys
import os
import glob
import spacy
import numpy as np
import fitz  # PyMuPDF
from typing import List, Dict, Tuple
from docx import Document as DocxDocument
from odf import text as odf_text
from odf.opendocument import load as load_odt

# Try to import tqdm for progress bars
try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False
    # Fallback: simple pass-through iterator
    def tqdm(iterable, **kwargs):
        return iterable


# ── Windows PyInstaller: hide console on GUI launch ───────────────────────────
# When launched via double-click or with --gui, minimize the console window
# so it doesn't appear in front of the GUI. The console keeps ownership of the
# process (required for correct I/O when launched from a terminal).
if sys.platform == "win32" and getattr(sys, "frozen", False):
    if len(sys.argv) == 1 or "--gui" in sys.argv:
        import ctypes
        _hwnd = ctypes.windll.kernel32.GetConsoleWindow()
        if _hwnd:
            ctypes.windll.user32.ShowWindow(_hwnd, 6)  # SW_MINIMIZE = 6
# ─────────────────────────────────────────────────────────────────────────────

APP_VERSION = "2.0"
APP_AUTHOR  = "Andrea Ciarrocchi"
APP_REPO    = "https://github.com/andreaciarrocchi/lexifinder"
APP_WEB     = "https://andreaciarrocchi.altervista.org"
APP_PAYPAL  = "https://paypal.me/ciarro85"


class DocumentReader:
    """Handles reading text from different document formats."""
    
    @staticmethod
    def read_pdf_text(filepath: str) -> str:
        """Extract full text from PDF."""
        doc = fitz.open(filepath)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        return text
    
    @staticmethod
    def read_pdf_by_page(filepath: str) -> List[str]:
        """Extract text page by page from PDF."""
        doc = fitz.open(filepath)
        pages = []
        for page in doc:
            pages.append(page.get_text())
        return pages
    
    @staticmethod
    def read_pdf_with_structure(filepath: str) -> List[Tuple[str, int, int]]:
        """
        Extract paragraphs from PDF with page and paragraph-in-page info.
        
        Returns:
            List of tuples: (paragraph_text, page_number, paragraph_in_page)
        """
        doc = fitz.open(filepath)
        structured_paragraphs = []
        
        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            # Split into paragraphs (by double newline or single newline)
            paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
            
            for para_in_page, para_text in enumerate(paragraphs, 1):
                structured_paragraphs.append((para_text, page_num, para_in_page))
        
        return structured_paragraphs
    
    @staticmethod
    def read_docx_text(filepath: str) -> str:
        """Extract full text from DOCX."""
        doc = DocxDocument(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    
    @staticmethod
    def read_docx_with_structure(filepath: str) -> List[Tuple[str, int, int]]:
        """
        Extract paragraphs from DOCX with chapter structure.
        
        Returns:
            List of tuples: (paragraph_text, chapter_number, paragraph_in_chapter)
        """
        doc = DocxDocument(filepath)
        structured_paragraphs = []
        current_chapter = 1
        para_in_chapter = 0
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
            
            # Check if this is a heading (new chapter/section)
            is_heading = False
            if paragraph.style.name.startswith('Heading'):
                # Heading 1 and Heading 2 are considered chapter markers
                heading_level = paragraph.style.name.replace('Heading ', '').strip()
                if heading_level in ['1', '2']:
                    is_heading = True
                    current_chapter += 1
                    para_in_chapter = 0
            
            # Don't index the heading itself, just use it as a marker
            if not is_heading:
                para_in_chapter += 1
                structured_paragraphs.append((paragraph.text, current_chapter, para_in_chapter))
        
        return structured_paragraphs
    
    @staticmethod
    def read_odt_text(filepath: str) -> str:
        """Extract full text from ODT."""
        doc = load_odt(filepath)
        paragraphs = []
        for paragraph in doc.getElementsByType(odf_text.P):
            text = str(paragraph)
            # Extract text content, removing XML tags
            import re
            text = re.sub('<[^>]+>', '', text)
            if text.strip():
                paragraphs.append(text.strip())
        return "\n".join(paragraphs)
    
    @staticmethod
    def read_odt_with_structure(filepath: str) -> List[Tuple[str, int, int]]:
        """
        Extract paragraphs from ODT with chapter structure.
        
        Returns:
            List of tuples: (paragraph_text, chapter_number, paragraph_in_chapter)
        """
        doc = load_odt(filepath)
        structured_paragraphs = []
        current_chapter = 1
        para_in_chapter = 0
        
        import re
        from odf import text as odf_text
        
        for element in doc.getElementsByType(odf_text.P) + doc.getElementsByType(odf_text.H):
            text_content = str(element)
            text_content = re.sub('<[^>]+>', '', text_content).strip()
            
            if not text_content:
                continue
            
            # Check if this is a heading
            is_heading = element.qname[1] == 'h'  # odf_text.H
            
            if is_heading:
                current_chapter += 1
                para_in_chapter = 0
            else:
                para_in_chapter += 1
                structured_paragraphs.append((text_content, current_chapter, para_in_chapter))
        
        return structured_paragraphs


class LexiFinderCLI:
    # Default and supported models
    DEFAULT_MODEL = "en_core_web_md"
    SUPPORTED_MODELS = {
        "en_core_web_sm": "English (small)",
        "en_core_web_md": "English (medium)",
        "en_core_web_lg": "English (large)",
        "it_core_news_sm": "Italian (small)",
        "it_core_news_md": "Italian (medium)",
        "it_core_news_lg": "Italian (large)",
        "de_core_news_sm": "German (small)",
        "de_core_news_md": "German (medium)",
        "de_core_news_lg": "German (large)",
        "fr_core_news_sm": "French (small)",
        "fr_core_news_md": "French (medium)",
        "fr_core_news_lg": "French (large)",
        "es_core_news_sm": "Spanish (small)",
        "es_core_news_md": "Spanish (medium)",
        "es_core_news_lg": "Spanish (large)",
    }
    
    # Indexing strategies
    STRATEGIES = {
        "keywords": "User-provided keywords (default - 2 levels)",
        "auto": "Automatic clustering (data-driven categories)",
        "hybrid": "Keywords with automatic sub-clustering (3 levels)",
        "frequent": "Most frequent terms as main categories"
    }

    # Chapter prefix per language (ISO 639-1 code → abbreviation)
    CHAPTER_PREFIXES = {
        "en": "Ch. ",   # English  – Chapter
        "it": "Ca. ",   # Italian  – Capitolo
        "de": "Kap. ",  # German   – Kapitel
        "fr": "Ch. ",   # French   – Chapitre
        "es": "Cap. ",  # Spanish  – Capítulo
    }
    DEFAULT_CHAPTER_PREFIX = "Ch. "  # fallback for unknown languages
    
    # Generic words to exclude (common across languages)
    GENERIC_WORDS = {
        # English
        'thing', 'way', 'part', 'aspect', 'factor', 'element', 'item', 'point',
        'type', 'kind', 'form', 'level', 'area', 'place', 'time', 'case',
        'example', 'instance', 'use', 'number', 'amount', 'value', 'result',
        'effect', 'process', 'system', 'method', 'approach', 'technique',
        'issue', 'problem', 'question', 'matter', 'situation', 'condition',
        'state', 'status', 'position', 'location', 'context', 'background',
        'overview', 'summary', 'conclusion', 'introduction', 'section', 'chapter',
        'figure', 'table', 'page', 'line', 'paragraph', 'sentence', 'word',
        # Italian
        'cosa', 'modo', 'parte', 'aspetto', 'fattore', 'elemento', 'tipo',
        'forma', 'livello', 'area', 'luogo', 'tempo', 'caso', 'esempio',
        'uso', 'numero', 'valore', 'risultato', 'effetto', 'processo',
        # German
        'ding', 'weise', 'teil', 'aspekt', 'faktor', 'element', 'typ',
        'form', 'ebene', 'bereich', 'ort', 'zeit', 'fall', 'beispiel',
        # French
        'chose', 'manière', 'partie', 'aspect', 'facteur', 'élément', 'type',
        'forme', 'niveau', 'zone', 'lieu', 'temps', 'cas', 'exemple',
        # Spanish
        'cosa', 'manera', 'parte', 'aspecto', 'factor', 'elemento', 'tipo',
        'forma', 'nivel', 'área', 'lugar', 'tiempo', 'caso', 'ejemplo',
    }
    
    # Default configuration
    DEFAULT_STRATEGY = "keywords"
    DEFAULT_MAX_PER_CATEGORY = 30
    DEFAULT_CLUSTERS = 8
    DEFAULT_SUBCLUSTERS = 3
    DEFAULT_MIN_OCCURRENCES = 1
    
    @staticmethod
    def get_models_directory():
        """Get the directory where external models are stored."""
        if getattr(sys, 'frozen', False):
            # Running as compiled executable
            app_dir = os.path.dirname(sys.executable)
        else:
            # Running as script
            app_dir = os.path.expanduser("~")
        
        models_dir = os.path.join(app_dir, ".lexifinder_models")
        os.makedirs(models_dir, exist_ok=True)
        return models_dir
    
    @staticmethod
    def download_model(model_name: str) -> bool:
        """
        Download a spaCy model.
        
        Args:
            model_name: Name of the model to download
            
        Returns:
            True if successful, False otherwise
        """
        import subprocess
        
        print(f"\nDownloading spaCy model '{model_name}'...")
        print("This may take a few minutes depending on the model size.\n")
        
        try:
            result = subprocess.run(
                [sys.executable, "-m", "spacy", "download", model_name],
                stderr=subprocess.PIPE,
                text=True
            )
            
            if result.returncode == 0:
                print(f"✓ Model '{model_name}' downloaded successfully!")
                return True
            else:
                print(f"✗ Error downloading model: {result.stderr}")
                return False
        except Exception as e:
            print(f"✗ Error: {e}")
            return False
    
    @staticmethod
    def delete_model(model_name: str, force: bool = False) -> bool:
        """
        Delete a spaCy model.
        
        Args:
            model_name: Name of the model to delete
            force: If True, skip confirmation prompt
            
        Returns:
            True if successful, False otherwise
        """
        import subprocess
        import shutil
        
        # Prevent deletion of bundled default model
        if model_name == LexiFinderCLI.DEFAULT_MODEL and getattr(sys, 'frozen', False):
            print(f"\n✗ Cannot delete '{model_name}' - it is the bundled default model.")
            print("  This model is integrated into the executable and cannot be removed.")
            return False
        
        # Check if model exists
        try:
            __import__(model_name)
            model_exists = True
        except ImportError:
            model_exists = False
        
        if not model_exists:
            print(f"\n✗ Model '{model_name}' is not installed.")
            return False
        
        # Confirmation prompt
        if not force:
            print(f"\n⚠ WARNING: You are about to delete model '{model_name}'")
            response = input("Are you sure? Type 'yes' to confirm: ")
            if response.lower() != 'yes':
                print("Deletion cancelled.")
                return False
        
        print(f"\nDeleting model '{model_name}'...")
        
        # Try to uninstall via pip
        try:
            result = subprocess.run(
                [sys.executable, "-m", "pip", "uninstall", "-y", model_name],
                capture_output=True,
                text=True
            )
            
            if result.returncode == 0:
                print(f"✓ Model '{model_name}' deleted successfully!")
                return True
            else:
                # If pip uninstall fails, try to remove manually
                print(f"⚠ pip uninstall failed, attempting manual removal...")
                
                # Try to find and remove the model directory
                try:
                    import site
                    import glob
                    
                    # Check in site-packages
                    for site_dir in site.getsitepackages():
                        model_path = os.path.join(site_dir, model_name)
                        if os.path.exists(model_path):
                            shutil.rmtree(model_path)
                            print(f"✓ Removed model directory: {model_path}")
                        
                        # Also check for .dist-info
                        for dist_info in glob.glob(os.path.join(site_dir, f"{model_name}*.dist-info")):
                            shutil.rmtree(dist_info)
                            print(f"✓ Removed metadata: {dist_info}")
                    
                    print(f"✓ Model '{model_name}' deleted successfully!")
                    return True
                    
                except Exception as e:
                    print(f"✗ Error during manual removal: {e}")
                    return False
                
        except Exception as e:
            print(f"✗ Error deleting model: {e}")
            return False
    
    @staticmethod
    def list_installed_models():
        """List all currently installed spaCy models."""
        print("\n" + "="*70)
        print("INSTALLED SPACY MODELS")
        print("="*70)
        
        installed = []
        
        # Check for known models
        for model_name in LexiFinderCLI.SUPPORTED_MODELS.keys():
            try:
                __import__(model_name)
                installed.append(model_name)
            except ImportError:
                pass
        
        if not installed:
            print("  No models installed.")
        else:
            for model_name in sorted(installed):
                desc = LexiFinderCLI.SUPPORTED_MODELS.get(model_name, "Unknown")
                
                # Check if it's the default bundled model
                is_bundled = (model_name == LexiFinderCLI.DEFAULT_MODEL and 
                             getattr(sys, 'frozen', False))
                
                if is_bundled:
                    print(f"  {model_name:25} - {desc} [BUNDLED]")
                else:
                    print(f"  {model_name:25} - {desc}")
        
        print("="*70)
        print(f"\nTotal installed: {len(installed)}")
        
        if installed:
            print("\nTo delete a model:")
            print("  lexifinder --delete-model <model_name>")
        print("")
    
    @staticmethod
    def save_config(config_path: str, args):
        """
        Save current configuration to a JSON file.
        
        Args:
            config_path: Path to save configuration file
            args: Parsed arguments from argparse
        """
        import json
        
        # Build configuration dictionary
        config = {}
        
        # Core parameters
        if args.keywords:
            config['keywords'] = args.keywords
        config['threshold'] = args.threshold
        config['mode'] = args.mode
        
        # Strategy parameters
        config['strategy'] = args.strategy
        config['max_per_category'] = args.max_per_category
        
        if args.strategy == 'auto':
            config['clusters'] = args.clusters
        elif args.strategy == 'hybrid':
            config['subclusters'] = args.subclusters
        elif args.strategy == 'frequent':
            config['top'] = args.top
        
        # Filtering parameters
        if args.exclude_generic:
            config['exclude_generic'] = True
        if args.min_occurrences > 1:
            config['min_occurrences'] = args.min_occurrences
        
        # Model parameter
        if args.model:
            config['model'] = args.model
        
        # Export format
        if args.export_format and args.export_format != 'txt':
            config['export_format'] = args.export_format
        
        # Document marking
        if args.mark:
            config['mark'] = True
        
        # Save to file
        try:
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=2, ensure_ascii=False)
            print(f"\n✓ Configuration saved to '{config_path}'")
            print(f"  Use with: lexifinder -i file.pdf -o output.txt --load-config {config_path}")
        except Exception as e:
            print(f"✗ Error saving configuration: {e}")
            sys.exit(1)
    
    @staticmethod
    def load_config(config_path: str) -> dict:
        """
        Load configuration from a JSON file.
        
        Args:
            config_path: Path to configuration file
            
        Returns:
            Dictionary with configuration parameters
        """
        import json
        
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
            print(f"✓ Configuration loaded from '{config_path}'")
            return config
        except FileNotFoundError:
            print(f"✗ Error: Configuration file '{config_path}' not found")
            sys.exit(1)
        except json.JSONDecodeError as e:
            print(f"✗ Error: Invalid JSON in configuration file: {e}")
            sys.exit(1)
        except Exception as e:
            print(f"✗ Error loading configuration: {e}")
            sys.exit(1)
    
    @staticmethod
    def cluster_nouns_by_similarity(nouns: List[str], nlp, n_clusters: int = 8) -> Dict[str, List[str]]:
        """
        Cluster nouns using K-means on word vectors.
        
        Args:
            nouns: List of nouns to cluster
            nlp: spaCy model
            n_clusters: Number of clusters to create
            
        Returns:
            Dictionary mapping cluster names to lists of nouns
        """
        import numpy as np
        
        # Get vectors for all nouns via direct vocab lookup (no NLP pipeline)
        vectors = []
        valid_nouns = []
        
        for noun in nouns:
            lex = nlp.vocab[noun]
            if lex.has_vector:
                vectors.append(lex.vector)
                valid_nouns.append(noun)
        
        if len(valid_nouns) < n_clusters:
            # Not enough nouns, create single cluster
            return {"Group 1": valid_nouns}
        
        # Simple K-means clustering using numpy
        vectors = np.array(vectors)
        
        # Initialize centroids randomly
        np.random.seed(42)
        indices = np.random.choice(len(vectors), n_clusters, replace=False)
        centroids = vectors[indices]
        
        # K-means iterations
        for _ in range(10):
            # Assign to nearest centroid
            distances = np.linalg.norm(vectors[:, np.newaxis] - centroids, axis=2)
            labels = np.argmin(distances, axis=1)
            
            # Update centroids
            new_centroids = np.array([vectors[labels == i].mean(axis=0) 
                                     if np.any(labels == i) else centroids[i]
                                     for i in range(n_clusters)])
            
            if np.allclose(centroids, new_centroids):
                break
            centroids = new_centroids
        
        # Group nouns by cluster
        clusters = {}
        for i in range(n_clusters):
            cluster_nouns = [valid_nouns[j] for j in range(len(valid_nouns)) if labels[j] == i]
            if cluster_nouns:
                # Name cluster based on most central noun
                cluster_center = centroids[i]
                distances_to_center = [np.linalg.norm(vectors[j] - cluster_center) 
                                      for j in range(len(valid_nouns)) if labels[j] == i]
                most_central_idx = np.argmin(distances_to_center)
                central_noun = cluster_nouns[most_central_idx]
                
                cluster_name = f"{central_noun.capitalize()} & Related"
                clusters[cluster_name] = sorted(cluster_nouns)
        
        return clusters
    
    @staticmethod
    def get_most_frequent_nouns(nouns: List[str], text: str, top_n: int = 15) -> List[str]:
        """
        Get the most frequent nouns from a document.
        
        Args:
            nouns: List of all nouns
            text: Document text
            top_n: Number of top nouns to return
            
        Returns:
            List of most frequent nouns
        """
        text_lower = text.lower()
        noun_counts = {}
        
        for noun in nouns:
            count = text_lower.count(noun.lower())
            if count > 0:
                noun_counts[noun] = count
        
        # Sort by frequency
        sorted_nouns = sorted(noun_counts.items(), key=lambda x: x[1], reverse=True)
        return [noun for noun, _ in sorted_nouns[:top_n]]
    
    @staticmethod
    def apply_smart_filters(noun_dict: Dict[str, List[str]], text: str = "", 
                           exclude_generic: bool = False, 
                           min_occurrences: int = 1) -> Dict[str, List[str]]:
        """
        Apply smart filtering to remove generic terms and low-frequency terms.
        
        Args:
            noun_dict: Dictionary mapping categories to noun lists
            text: Full document text (for occurrence counting)
            exclude_generic: If True, exclude generic/common words
            min_occurrences: Minimum occurrences required to keep a term
            
        Returns:
            Filtered dictionary
        """
        filtered = {}
        text_lower = text.lower() if text else ""
        
        for category, nouns in noun_dict.items():
            filtered_nouns = []
            
            for noun in nouns:
                # Check generic word exclusion
                if exclude_generic and noun.lower() in LexiFinderCLI.GENERIC_WORDS:
                    continue
                
                # Check minimum occurrences
                if min_occurrences > 1 and text:
                    count = text_lower.count(noun.lower())
                    if count < min_occurrences:
                        continue
                
                filtered_nouns.append(noun)
            
            if filtered_nouns:  # Only keep categories with remaining nouns
                filtered[category] = filtered_nouns
        
        return filtered
    
    @staticmethod
    def apply_smart_filters_3level(noun_dict: Dict[str, Dict[str, List[str]]], 
                                   text: str = "", 
                                   exclude_generic: bool = False, 
                                   min_occurrences: int = 1) -> Dict[str, Dict[str, List[str]]]:
        """
        Apply smart filtering to 3-level hierarchy.
        
        Args:
            noun_dict: 3-level dictionary (main -> sub -> nouns)
            text: Full document text
            exclude_generic: If True, exclude generic words
            min_occurrences: Minimum occurrences required
            
        Returns:
            Filtered 3-level dictionary
        """
        filtered = {}
        text_lower = text.lower() if text else ""
        
        for main_cat, sub_dict in noun_dict.items():
            filtered_sub = {}
            
            for sub_cat, nouns in sub_dict.items():
                filtered_nouns = []
                
                for noun in nouns:
                    # Check generic word exclusion
                    if exclude_generic and noun.lower() in LexiFinderCLI.GENERIC_WORDS:
                        continue
                    
                    # Check minimum occurrences
                    if min_occurrences > 1 and text:
                        count = text_lower.count(noun.lower())
                        if count < min_occurrences:
                            continue
                    
                    filtered_nouns.append(noun)
                
                if filtered_nouns:
                    filtered_sub[sub_cat] = filtered_nouns
            
            if filtered_sub:
                filtered[main_cat] = filtered_sub
        
        return filtered
    
    @staticmethod
    def list_available_models():
        """List all available spaCy models."""
        print("\n" + "="*70)
        print("AVAILABLE SPACY MODELS")
        print("="*70)
        for model, desc in LexiFinderCLI.SUPPORTED_MODELS.items():
            print(f"  {model:25} - {desc}")
        print("="*70)
        print("\nTo download a model:")
        print("  lexifinder --download-model <model_name>")
        print("\nTo see installed models:")
        print("  lexifinder --list-installed")
        print("\nTo delete a model:")
        print("  lexifinder --delete-model <model_name>")
        print("")
    
    @staticmethod
    def _resolve_model_path(model_name: str) -> str:
        """
        Resolve the spaCy model name to a filesystem path when running as a
        PyInstaller frozen executable.

        When frozen, the default bundled model (en_core_web_md) is unpacked
        by PyInstaller into sys._MEIPASS.  If we pass the bare model *name*
        to spacy.load() in that context it fails because the model is not
        registered as an installed Python package.  We therefore substitute
        the absolute path so spaCy can find it directly.

        For non-frozen execution (normal Python interpreter) the model name
        is returned unchanged, preserving the original behaviour.

        Args:
            model_name: spaCy model name (e.g. "en_core_web_md")

        Returns:
            Absolute path string (frozen) or the original model name (script).
        """
        if getattr(sys, 'frozen', False):
            # sys._MEIPASS is the temp folder where PyInstaller unpacks data
            bundled_path = os.path.join(sys._MEIPASS, model_name)
            if os.path.isdir(bundled_path):
                return bundled_path
        return model_name

    def __init__(self, similarity_threshold: float = 0.5, model_name: str = None):
        """
        Initialize the LexiFinder with a spaCy language model.
        
        Args:
            similarity_threshold: Minimum similarity score (0.0 to 1.0)
            model_name: Name of the spaCy model to use (default: en_core_web_md)
        """
        if model_name is None:
            model_name = self.DEFAULT_MODEL
        
        print(f"Loading spaCy language model '{model_name}'...")
        
        # Try to load the specified model
        try:
            self.nlp = spacy.load(self._resolve_model_path(model_name))
            print(f"✓ Model '{model_name}' loaded successfully.")
        except OSError:
            # Model not found
            if model_name == self.DEFAULT_MODEL:
                # Default model should be bundled
                print(f"✗ Error: Default model '{self.DEFAULT_MODEL}' not found.")
                print("This should not happen in a properly packaged version.")
                print(f"Please reinstall or download: python -m spacy download {self.DEFAULT_MODEL}")
                sys.exit(1)
            else:
                # Non-default model not found - offer to download
                print(f"✗ Model '{model_name}' not found.")
                
                # Check if it's a known model
                if model_name in self.SUPPORTED_MODELS:
                    response = input(f"\nWould you like to download '{model_name}' now? [y/N]: ")
                    if response.lower() in ['y', 'yes']:
                        if self.download_model(model_name):
                            try:
                                self.nlp = spacy.load(self._resolve_model_path(model_name))
                                print(f"✓ Model '{model_name}' loaded successfully.")
                            except:
                                print(f"✗ Failed to load model after download.")
                                print(f"Falling back to default model '{self.DEFAULT_MODEL}'...")
                                self._load_fallback_model()
                        else:
                            print(f"Falling back to default model '{self.DEFAULT_MODEL}'...")
                            self._load_fallback_model()
                    else:
                        print(f"Falling back to default model '{self.DEFAULT_MODEL}'...")
                        self._load_fallback_model()
                else:
                    print(f"Unknown model '{model_name}'.")
                    print(f"Falling back to default model '{self.DEFAULT_MODEL}'...")
                    self._load_fallback_model()
        
        self.similarity_threshold = similarity_threshold
        self.reader = DocumentReader()
        self.chapter_prefix = self.CHAPTER_PREFIXES.get(
            self.nlp.lang, self.DEFAULT_CHAPTER_PREFIX
        )
    
    def _load_fallback_model(self):
        """Load the default fallback model."""
        try:
            self.nlp = spacy.load(self._resolve_model_path(self.DEFAULT_MODEL))
            print(f"✓ Fallback model '{self.DEFAULT_MODEL}' loaded successfully.")
        except OSError:
            print(f"✗ Critical error: Cannot load fallback model '{self.DEFAULT_MODEL}'.")
            print("Please ensure spaCy and the model are properly installed.")
            sys.exit(1)

    def detect_file_type(self, filepath: str) -> str:
        """
        Detect document type from extension.
        
        Args:
            filepath: Path to the document
            
        Returns:
            File type: 'pdf', 'docx', or 'odt'
        """
        ext = filepath.lower()
        if ext.endswith('.pdf'):
            return 'pdf'
        elif ext.endswith('.docx'):
            return 'docx'
        elif ext.endswith('.odt'):
            return 'odt'
        else:
            raise ValueError(f"Unsupported file format. Supported: .pdf, .docx, .odt")

    def extract_nouns(self, filepath: str) -> List[str]:
        """
        Extract unique nouns from a document.
        
        Args:
            filepath: Path to the document file
            
        Returns:
            Sorted list of unique lemmatized nouns
        """
        file_type = self.detect_file_type(filepath)
        print(f"Extracting nouns from '{filepath}' ({file_type.upper()})...")
        
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        # Read full text based on file type
        if file_type == 'pdf':
            text = self.reader.read_pdf_text(filepath)
        elif file_type == 'docx':
            text = self.reader.read_docx_text(filepath)
        elif file_type == 'odt':
            text = self.reader.read_odt_text(filepath)

        # Split the text into chunks so that nlp.pipe() can report
        # real incremental progress via tqdm instead of blocking silently
        # on a single large nlp(text) call (which gives zero feedback and
        # can take minutes on large documents).
        _CHUNK_CHARS = 25_000          # ~25 k chars per chunk (~3–5 k words)
        _chunks = []
        for _i in range(0, max(1, len(text)), _CHUNK_CHARS):
            _chunk = text[_i:_i + _CHUNK_CHARS].strip()
            if _chunk:
                _chunks.append(_chunk)
        if not _chunks:
            _chunks = [text]

        n_chunks = len(_chunks)
        print(f"  Analyzing text with spaCy ({n_chunks} chunk{'s' if n_chunks != 1 else ''})...")

        unique_nouns = set()
        for _doc in tqdm(
            self.nlp.pipe(_chunks, batch_size=1),
            desc="  Analyzing text",
            total=n_chunks,
            unit="chunk",
            disable=not TQDM_AVAILABLE,
            ncols=80,
        ):
            for token in _doc:
                if token.pos_ == "NOUN" and not token.is_punct and not token.is_space:
                    unique_nouns.add(token.lemma_.lower())

        print(f"  Found {len(unique_nouns)} unique nouns.")
        return sorted(list(unique_nouns))

    def find_correlated_nouns(self, nouns: List[str], keywords: List[str] = None, 
                             text: str = "", strategy: str = "keywords", 
                             **options) -> Dict[str, any]:
        """
        Create hierarchical structure using the specified strategy.
        
        Args:
            nouns: List of nouns to organize
            keywords: List of reference keywords (for keywords/hybrid strategies)
            text: Full document text (for frequent strategy and filtering)
            strategy: Indexing strategy ('keywords', 'auto', 'hybrid', 'frequent')
            **options: Additional options (max_per_category, clusters, subclusters, 
                      exclude_generic, min_occurrences)
            
        Returns:
            Hierarchical structure (format depends on strategy)
        """
        # Get strategy result
        if strategy == "keywords":
            result = self._strategy_keywords(nouns, keywords, options)
        elif strategy == "auto":
            result = self._strategy_auto(nouns, options)
        elif strategy == "hybrid":
            result = self._strategy_hybrid(nouns, keywords, options)
        elif strategy == "frequent":
            result = self._strategy_frequent(nouns, text, options)
        else:
            print(f"Warning: Unknown strategy '{strategy}', using 'keywords'")
            result = self._strategy_keywords(nouns, keywords, options)
        
        # Apply smart filtering if requested
        exclude_generic = options.get('exclude_generic', False)
        min_occurrences = options.get('min_occurrences', self.DEFAULT_MIN_OCCURRENCES)
        
        if exclude_generic or min_occurrences > 1:
            print(f"\nApplying smart filters...")
            if exclude_generic:
                print(f"  Excluding generic/common words")
            if min_occurrences > 1:
                print(f"  Excluding terms with < {min_occurrences} occurrences")
            
            # Check if 3-level (hybrid) or 2-level
            if result and isinstance(next(iter(result.values())), dict):
                # 3-level hierarchy
                result = self.apply_smart_filters_3level(result, text, exclude_generic, min_occurrences)
            else:
                # 2-level hierarchy
                result = self.apply_smart_filters(result, text, exclude_generic, min_occurrences)
            
            # Count remaining terms
            if result and isinstance(next(iter(result.values())), dict):
                total = sum(sum(len(v) for v in cat.values()) for cat in result.values())
            else:
                total = sum(len(v) for v in result.values())
            print(f"  {total} terms remaining after filtering")
        
        return result
    
    @staticmethod
    def _vocab_vector(nlp, word: str):
        """
        Return (vector, has_vector) for *word* by looking it up directly in
        spaCy's vocabulary — WITHOUT running any NLP pipeline stage.

        This is the critical performance fix.  The previous implementation
        called nlp.pipe(nouns) which runs the full tokenizer + tagger +
        parser + NER pipeline on every single noun, creates a heavyweight
        Cython Doc object for each one, and then stores thousands of those
        Doc objects in a dict.  When the dict is eventually garbage-collected
        Python's reference-counter has to deallocate every Doc (each of which
        holds POS tags, dependency arcs, entity spans, etc.) — and for
        thousands of items this GC pause can last several minutes while
        holding the GIL, freezing the Flet GUI completely.

        nlp.vocab[word] is an O(1) hash-table lookup that returns a Lexeme —
        a tiny, immutable vocabulary record that holds just the word's
        pre-computed vector as a numpy array.  No pipeline, no Doc, no GC.
        """
        lex = nlp.vocab[word]
        return lex.vector, lex.has_vector

    @staticmethod
    def _cosine_sim(v1, v2, n1: float, n2: float) -> float:
        """
        Cosine similarity between two pre-normalised vectors.
        Pass the norms pre-computed to avoid redundant sqrt in inner loops.
        Returns 0.0 when either vector is the zero vector.
        """
        if n1 == 0.0 or n2 == 0.0:
            return 0.0
        return float(np.dot(v1, v2) / (n1 * n2))

    def _strategy_keywords(self, nouns: List[str], keywords: List[str], 
                          options: dict) -> Dict[str, List[str]]:
        """
        Original strategy: User keywords → Similar nouns (2 levels)
        """
        print(f"Building index using KEYWORDS strategy (threshold: {self.similarity_threshold})...")
        
        max_per_category = options.get('max_per_category', self.DEFAULT_MAX_PER_CATEGORY)
        hierarchical_index = {keyword: [] for keyword in keywords}

        # Keyword vectors — vocab lookup, instant, no Doc objects created
        keyword_vecs = []
        for kw in keywords:
            vec, has_vec = self._vocab_vector(self.nlp, kw)
            norm = float(np.linalg.norm(vec))
            keyword_vecs.append((kw, vec, norm, has_vec))

        # Noun vectors — vocab lookup, one entry per noun, zero NLP pipeline
        print(f"  Pre-computing vectors for {len(nouns)} nouns...")
        noun_vecs = []
        for noun in tqdm(
            nouns,
            desc="  Computing vectors",
            total=len(nouns),
            unit="noun",
            disable=not TQDM_AVAILABLE,
            ncols=80,
        ):
            vec, has_vec = self._vocab_vector(self.nlp, noun)
            norm = float(np.linalg.norm(vec))
            noun_vecs.append((noun, vec, norm, has_vec))

        total_matches = 0

        # Cosine-similarity comparison — pure numpy, no Doc objects
        for noun, n_vec, n_norm, n_has in tqdm(
            noun_vecs,
            desc="  Finding correlations",
            unit="noun",
            disable=not TQDM_AVAILABLE,
            ncols=80,
        ):
            if not n_has:
                continue
            for kw, k_vec, k_norm, k_has in keyword_vecs:
                if not k_has:
                    continue
                sim = self._cosine_sim(n_vec, k_vec, n_norm, k_norm)
                if sim >= self.similarity_threshold:
                    hierarchical_index[kw].append(noun)
                    total_matches += 1

        for keyword in hierarchical_index:
            hierarchical_index[keyword] = sorted(hierarchical_index[keyword])[:max_per_category]
            print(f"  '{keyword}': {len(hierarchical_index[keyword])} sub-entries")

        print(f"  Total sub-entries: {total_matches} (limited to {max_per_category} per category)")
        return hierarchical_index
    
    def _strategy_auto(self, nouns: List[str], options: dict) -> Dict[str, List[str]]:
        """
        Automatic clustering: Data-driven categories (2 levels)
        """
        print(f"Building index using AUTO-CLUSTERING strategy...")
        
        n_clusters = options.get('clusters', self.DEFAULT_CLUSTERS)
        max_per_category = options.get('max_per_category', self.DEFAULT_MAX_PER_CATEGORY)
        
        print(f"  Creating {n_clusters} clusters...")
        clusters = self.cluster_nouns_by_similarity(nouns, self.nlp, n_clusters)
        
        # Limit entries per cluster
        for cluster_name in clusters:
            clusters[cluster_name] = clusters[cluster_name][:max_per_category]
            print(f"  '{cluster_name}': {len(clusters[cluster_name])} sub-entries")
        
        total = sum(len(v) for v in clusters.values())
        print(f"  Total sub-entries: {total}")
        return clusters
    
    def _strategy_hybrid(self, nouns: List[str], keywords: List[str], 
                        options: dict) -> Dict[str, Dict[str, List[str]]]:
        """
        Hybrid strategy: Keywords → Sub-clusters → Nouns (3 levels)
        """
        print(f"Building index using HYBRID strategy (keywords + sub-clustering)...")
        
        subclusters = options.get('subclusters', self.DEFAULT_SUBCLUSTERS)
        max_per_category = options.get('max_per_category', self.DEFAULT_MAX_PER_CATEGORY)
        hierarchical_index = {}

        # Keyword vectors from vocab — no NLP pipeline, no Doc objects
        keyword_vecs = []
        for kw in keywords:
            vec, has_vec = self._vocab_vector(self.nlp, kw)
            norm = float(np.linalg.norm(vec))
            keyword_vecs.append((kw, vec, norm, has_vec))

        # Noun vectors from vocab — instant O(1) lookup per word
        print(f"  Pre-computing vectors for {len(nouns)} nouns...")
        noun_vecs = []
        for noun in tqdm(
            nouns,
            desc="  Computing vectors",
            total=len(nouns),
            unit="noun",
            disable=not TQDM_AVAILABLE,
            ncols=80,
        ):
            vec, has_vec = self._vocab_vector(self.nlp, noun)
            norm = float(np.linalg.norm(vec))
            noun_vecs.append((noun, vec, norm, has_vec))

        # For each keyword, find similar nouns using cached vectors
        for kw, k_vec, k_norm, k_has in tqdm(keyword_vecs, desc="  Processing keywords",
                                               unit="keyword", disable=not TQDM_AVAILABLE,
                                               ncols=80):
            similar_nouns = []
            if k_has and k_norm > 0:
                for noun, n_vec, n_norm, n_has in noun_vecs:
                    if n_has:
                        sim = self._cosine_sim(n_vec, k_vec, n_norm, k_norm)
                        if sim >= self.similarity_threshold:
                            similar_nouns.append(noun)

            if len(similar_nouns) > subclusters:
                print(f"  '{kw}': Clustering {len(similar_nouns)} nouns into {subclusters} sub-groups...")
                sub_clusters = self.cluster_nouns_by_similarity(similar_nouns, self.nlp, subclusters)
                for sub_name in sub_clusters:
                    sub_clusters[sub_name] = sorted(sub_clusters[sub_name])[:max_per_category//subclusters]
                hierarchical_index[kw] = sub_clusters
                total_in_kw = sum(len(v) for v in sub_clusters.values())
                print(f"  '{kw}': {total_in_kw} sub-entries in {len(sub_clusters)} groups")
            else:
                hierarchical_index[kw] = {"Main Group": sorted(similar_nouns)[:max_per_category]}
                print(f"  '{kw}': {len(similar_nouns)} sub-entries (single group)")

        total_entries = sum(sum(len(v) for v in cat.values()) for cat in hierarchical_index.values())
        print(f"  Total sub-entries: {total_entries}")
        return hierarchical_index
    
    def _strategy_frequent(self, nouns: List[str], text: str, 
                          options: dict) -> Dict[str, List[str]]:
        """
        Frequent terms strategy: Most frequent nouns as categories (2 levels)
        """
        print(f"Building index using FREQUENT TERMS strategy...")
        
        top_n = options.get('top', 15)
        max_per_category = options.get('max_per_category', self.DEFAULT_MAX_PER_CATEGORY)
        
        frequent_nouns = self.get_most_frequent_nouns(nouns, text, top_n)
        print(f"  Using {len(frequent_nouns)} most frequent nouns as categories")

        # Noun vectors from vocab — instant O(1) lookup, no Doc objects
        print(f"  Pre-computing vectors for {len(nouns)} nouns...")
        noun_vecs = []
        for noun in tqdm(
            nouns,
            desc="  Computing vectors",
            total=len(nouns),
            unit="noun",
            disable=not TQDM_AVAILABLE,
            ncols=80,
        ):
            vec, has_vec = self._vocab_vector(self.nlp, noun)
            norm = float(np.linalg.norm(vec))
            noun_vecs.append((noun, vec, norm, has_vec))

        # Frequent-noun vectors (subset of the above)
        noun_vec_map = {n: (v, norm, hv) for n, v, norm, hv in noun_vecs}
        freq_vecs = [(fn, *noun_vec_map[fn]) for fn in frequent_nouns if fn in noun_vec_map]

        hierarchical_index = {}
        for freq_noun, f_vec, f_norm, f_has in tqdm(
            freq_vecs,
            desc="  Finding correlations",
            unit="category",
            disable=not TQDM_AVAILABLE,
            ncols=80,
        ):
            similar_nouns = []
            if f_has and f_norm > 0:
                for noun, n_vec, n_norm, n_has in noun_vecs:
                    if noun == freq_noun or not n_has:
                        continue
                    sim = self._cosine_sim(n_vec, f_vec, n_norm, f_norm)
                    if sim >= self.similarity_threshold:
                        similar_nouns.append(noun)
            hierarchical_index[freq_noun] = sorted(similar_nouns)[:max_per_category]
            print(f"  '{freq_noun}': {len(hierarchical_index[freq_noun])} sub-entries")

        total = sum(len(v) for v in hierarchical_index.values())
        print(f"  Total sub-entries: {total}")
        return hierarchical_index

    def extract_occurrences(self, filepath: str, hierarchical_nouns: Dict[str, List[str]], 
                           mode: str) -> Tuple[Dict[str, Dict[str, List[str]]], str]:
        """
        Build a hierarchical index of word occurrences with structured references.
        
        Args:
            filepath: Path to the document file
            hierarchical_nouns: Dictionary mapping keywords to their correlated nouns
            mode: 'page' or 'paragraph'
            
        Returns:
            Tuple of (nested dictionary: keyword -> noun -> list of reference strings, format_type)
        """
        file_type = self.detect_file_type(filepath)
        
        # Determine indexing mode and format based on file type
        if file_type in ['docx', 'odt']:
            actual_mode = 'paragraph'
            format_type = 'chapter_para'  # "Cap. X, §Y"
            if mode == 'page':
                print(f"  Note: {file_type.upper()} files don't have pages. Using paragraph mode.")
        else:  # PDF
            actual_mode = mode
            if mode == 'page':
                format_type = 'page'  # "p.X"
            else:
                format_type = 'page_para'  # "p.X, §Y"
        
        print(f"Building index by {actual_mode} for all sub-entries...")
        
        # Read document with structure
        if file_type == 'pdf':
            if actual_mode == 'page':
                # Original page-only mode
                pages = self.reader.read_pdf_by_page(filepath)
                return self._extract_by_page_only(pages, hierarchical_nouns)
            else:  # paragraph mode for PDF
                structured_data = self.reader.read_pdf_with_structure(filepath)
                return self._extract_with_structure(structured_data, hierarchical_nouns, format_type)
        elif file_type == 'docx':
            structured_data = self.reader.read_docx_with_structure(filepath)
            return self._extract_with_structure(structured_data, hierarchical_nouns, format_type)
        elif file_type == 'odt':
            structured_data = self.reader.read_odt_with_structure(filepath)
            return self._extract_with_structure(structured_data, hierarchical_nouns, format_type)
    
    def _flatten_hierarchy(self, hierarchical_nouns: any) -> Dict[str, List[str]]:
        """
        Flatten 3-level hierarchy to 2-level for extraction.
        
        Args:
            hierarchical_nouns: Either Dict[str, List[str]] or Dict[str, Dict[str, List[str]]]
            
        Returns:
            Dict[str, List[str]] - Flattened 2-level structure
        """
        flattened = {}
        
        for main_key, value in hierarchical_nouns.items():
            if isinstance(value, dict):
                # 3-level: flatten sub-clusters
                all_nouns = []
                for sub_cluster_nouns in value.values():
                    all_nouns.extend(sub_cluster_nouns)
                flattened[main_key] = all_nouns
            else:
                # 2-level: use as-is
                flattened[main_key] = value
        
        return flattened
    
    def _extract_by_page_only(self, pages: List[str], hierarchical_nouns: any) -> Tuple[Dict[str, any], str]:
        """Extract occurrences by page number only (PDF page mode)."""
        # Flatten if needed
        flat_hierarchy = self._flatten_hierarchy(hierarchical_nouns)
        
        all_nouns = set()
        for nouns_list in flat_hierarchy.values():
            all_nouns.update(nouns_list)
        
        noun_pages = {noun.lower(): [] for noun in all_nouns}
        
        # Progress bar for page indexing
        for page_num, page_text in tqdm(enumerate(pages, 1), 
                                        desc="  Indexing pages",
                                        total=len(pages),
                                        unit="page",
                                        disable=not TQDM_AVAILABLE,
                                        ncols=80):
            text_lower = page_text.lower()
            for noun in all_nouns:
                if noun.lower() in text_lower:
                    noun_pages[noun.lower()].append(f"p.{page_num}")
        
        # Organize by keyword, preserving original structure
        hierarchical_index = {}
        for keyword, value in hierarchical_nouns.items():
            if isinstance(value, dict):
                # 3-level: preserve sub-clusters
                hierarchical_index[keyword] = {}
                for sub_name, nouns_list in value.items():
                    hierarchical_index[keyword][sub_name] = {}
                    for noun in nouns_list:
                        refs = noun_pages[noun.lower()]
                        if refs:
                            hierarchical_index[keyword][sub_name][noun] = sorted(set(refs), key=lambda x: int(x.split('.')[1]))
            else:
                # 2-level: simple structure
                hierarchical_index[keyword] = {}
                for noun in value:
                    refs = noun_pages[noun.lower()]
                    if refs:
                        hierarchical_index[keyword][noun] = sorted(set(refs), key=lambda x: int(x.split('.')[1]))
        
        print(f"  Indexing complete ({len(pages)} pages processed).")
        return hierarchical_index, 'page'
    
    def _extract_with_structure(self, structured_data: List[Tuple[str, int, int]], 
                                hierarchical_nouns: any, 
                                format_type: str) -> Tuple[Dict[str, any], str]:
        """
        Extract occurrences with chapter/page and paragraph structure.
        
        Args:
            structured_data: List of (text, section_num, para_in_section)
            hierarchical_nouns: Keywords and their nouns (2 or 3 levels)
            format_type: 'chapter_para' or 'page_para'
        """
        # Flatten if needed
        flat_hierarchy = self._flatten_hierarchy(hierarchical_nouns)
        
        all_nouns = set()
        for nouns_list in flat_hierarchy.values():
            all_nouns.update(nouns_list)
        
        noun_refs = {noun.lower(): [] for noun in all_nouns}
        
        # Progress bar for paragraph indexing
        unit_desc = "paragraphs"
        for para_text, section_num, para_in_section in tqdm(structured_data,
                                                             desc=f"  Indexing {unit_desc}",
                                                             unit="para",
                                                             disable=not TQDM_AVAILABLE,
                                                             ncols=80):
            text_lower = para_text.lower()
            
            for noun in all_nouns:
                if noun.lower() in text_lower:
                    # Create reference string based on format
                    if format_type == 'chapter_para':
                        ref = f"{self.chapter_prefix}{section_num}, §{para_in_section}"
                    else:  # page_para
                        ref = f"p.{section_num}, §{para_in_section}"
                    noun_refs[noun.lower()].append(ref)
        
        # Organize by keyword, preserving original structure
        hierarchical_index = {}
        for keyword, value in hierarchical_nouns.items():
            if isinstance(value, dict):
                # 3-level: preserve sub-clusters
                hierarchical_index[keyword] = {}
                for sub_name, nouns_list in value.items():
                    hierarchical_index[keyword][sub_name] = {}
                    for noun in nouns_list:
                        refs = noun_refs[noun.lower()]
                        if refs:
                            sorted_refs = self._sort_references(refs, format_type)
                            hierarchical_index[keyword][sub_name][noun] = self._compact_references(sorted_refs, format_type)
            else:
                # 2-level: simple structure
                hierarchical_index[keyword] = {}
                for noun in value:
                    refs = noun_refs[noun.lower()]
                    if refs:
                        sorted_refs = self._sort_references(refs, format_type)
                        hierarchical_index[keyword][noun] = self._compact_references(sorted_refs, format_type)
        
        print(f"  Indexing complete ({len(structured_data)} paragraphs processed).")
        return hierarchical_index, format_type
    
    def _sort_references(self, refs: List[str], format_type: str) -> List[str]:
        """Sort and deduplicate references."""
        unique_refs = list(set(refs))
        
        if format_type in ['chapter_para', 'page_para']:
            # Sort by chapter/page number, then paragraph number
            def sort_key(ref):
                parts = ref.replace(self.chapter_prefix, '').replace('p.', '').replace('§', '').split(', ')
                return (int(parts[0]), int(parts[1]))
            return sorted(unique_refs, key=sort_key)
        else:  # 'page'
            return sorted(unique_refs, key=lambda x: int(x.split('.')[1]))
    
    def _compact_references(self, refs: List[str], format_type: str) -> List[str]:
        """
        Compact references by grouping paragraphs from same chapter/page.
        
        Example: ["Ch. 2, §1", "Ch. 2, §3", "Ch. 2, §5", "Ch. 3, §2"]
        Becomes: ["Ch. 2, §1, 3, 5", "Ch. 3, §2"]
        (the actual prefix depends on the loaded model language)
        
        Args:
            refs: Sorted list of references
            format_type: 'page', 'page_para', or 'chapter_para'
            
        Returns:
            Compacted list of references
        """
        if format_type not in ['chapter_para', 'page_para']:
            # No compaction needed for simple page references
            return refs
        
        if not refs:
            return refs
        
        compacted = []
        current_section = None
        current_paragraphs = []
        
        for ref in refs:
            # Parse reference
            if self.chapter_prefix in ref:
                prefix = self.chapter_prefix
                parts = ref.replace(self.chapter_prefix, '').replace('§', '').split(', ')
            else:  # 'p.'
                prefix = 'p.'
                parts = ref.replace('p.', '').replace('§', '').split(', ')
            
            section = int(parts[0])
            paragraph = int(parts[1])
            
            if section != current_section:
                # New section - save previous group
                if current_section is not None:
                    compacted.append(self._format_compacted_ref(prefix, current_section, current_paragraphs))
                current_section = section
                current_paragraphs = [paragraph]
            else:
                # Same section - add paragraph
                current_paragraphs.append(paragraph)
        
        # Add last group
        if current_section is not None:
            compacted.append(self._format_compacted_ref(prefix, current_section, current_paragraphs))
        
        return compacted
    
    def _format_compacted_ref(self, prefix: str, section: int, paragraphs: List[int]) -> str:
        """
        Format a compacted reference.
        
        Args:
            prefix: chapter prefix (language-dependent, e.g. 'Ch. ', 'Ca. ', 'Kap. ') or 'p.'
            section: Section/chapter/page number
            paragraphs: List of paragraph numbers
            
        Returns:
            Formatted reference string
        """
        if len(paragraphs) == 1:
            return f"{prefix}{section}, §{paragraphs[0]}"
        else:
            para_str = ', '.join(str(p) for p in paragraphs)
            return f"{prefix}{section}, §{para_str}"

    def write_index(self, hierarchical_index: Dict[str, any], 
                   output_path: str, format_type: str, export_format: str = 'txt'):
        """
        Write the hierarchical index in specified format.
        Supports both 2-level and 3-level hierarchies.
        
        Args:
            hierarchical_index: Nested dictionary (2 or 3 levels)
            output_path: Path to output file
            format_type: 'page', 'page_para', or 'chapter_para'
            export_format: 'txt', 'json', 'csv', or 'html'
        """
        # Ensure correct file extension
        base_path = os.path.splitext(output_path)[0]
        output_path = f"{base_path}.{export_format}"
        
        print(f"Writing hierarchical index to '{output_path}' ({export_format.upper()} format)...")
        
        # Ensure output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Export in requested format
        if export_format == 'json':
            self._export_json(hierarchical_index, output_path, format_type)
        elif export_format == 'csv':
            self._export_csv(hierarchical_index, output_path, format_type)
        elif export_format == 'html':
            self._export_html(hierarchical_index, output_path, format_type)
        else:  # txt (default)
            self._export_txt(hierarchical_index, output_path, format_type)
        
        print(f"  Index written successfully.")
    
    def _export_txt(self, hierarchical_index: Dict[str, any], 
                    output_path: str, format_type: str):
        """Export index as plain text."""
        with open(output_path, 'w', encoding='utf-8') as f:
            # Write hierarchical structure (clean output, no decorative headers)
            for main_concept in sorted(hierarchical_index.keys()):
                f.write(f"{main_concept.upper()}\n")
                
                entries = hierarchical_index[main_concept]
                
                # Check if this is 3-level (has sub-clusters) or 2-level
                if entries and isinstance(next(iter(entries.values())), dict):
                    # 3-level hierarchy: Main → Sub-cluster → Nouns
                    for sub_cluster in sorted(entries.keys()):
                        f.write(f"  [{sub_cluster}]\n")
                        sub_entries = entries[sub_cluster]
                        
                        for sub_entry in sorted(sub_entries.keys()):
                            refs = sub_entries[sub_entry]
                            refs_str = '; '.join(refs)
                            f.write(f"    {sub_entry}: {refs_str}\n")
                        
                        f.write("\n")
                else:
                    # 2-level hierarchy: Main → Nouns
                    for sub_entry in sorted(entries.keys()):
                        refs = entries[sub_entry]
                        refs_str = '; '.join(refs)
                        f.write(f"    {sub_entry}: {refs_str}\n")
                
                f.write("\n")
    
    def _export_json(self, hierarchical_index: Dict[str, any], 
                     output_path: str, format_type: str):
        """Export index as JSON."""
        import json
        
        # Build structured JSON
        output_data = {
            "format": format_type,
            "index": {}
        }
        
        for main_concept in hierarchical_index:
            entries = hierarchical_index[main_concept]
            
            # Check if 3-level or 2-level
            if entries and isinstance(next(iter(entries.values())), dict):
                # 3-level
                output_data["index"][main_concept] = {}
                for sub_cluster, sub_entries in entries.items():
                    output_data["index"][main_concept][sub_cluster] = dict(sub_entries)
            else:
                # 2-level
                output_data["index"][main_concept] = dict(entries)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)
    
    def _export_csv(self, hierarchical_index: Dict[str, any], 
                    output_path: str, format_type: str):
        """Export index as CSV."""
        import csv
        
        with open(output_path, 'w', encoding='utf-8', newline='') as f:
            writer = csv.writer(f)
            
            # Check if 3-level or 2-level structure
            first_entries = next(iter(hierarchical_index.values()))
            is_3level = first_entries and isinstance(next(iter(first_entries.values())), dict)
            
            # Write header
            if is_3level:
                writer.writerow(['Category', 'Subcategory', 'Term', 'References'])
            else:
                writer.writerow(['Category', 'Term', 'References'])
            
            # Write data
            for main_concept in sorted(hierarchical_index.keys()):
                entries = hierarchical_index[main_concept]
                
                if is_3level:
                    # 3-level hierarchy
                    for sub_cluster in sorted(entries.keys()):
                        sub_entries = entries[sub_cluster]
                        for term in sorted(sub_entries.keys()):
                            refs = sub_entries[term]
                            refs_str = '; '.join(refs)
                            writer.writerow([main_concept, sub_cluster, term, refs_str])
                else:
                    # 2-level hierarchy
                    for term in sorted(entries.keys()):
                        refs = entries[term]
                        refs_str = '; '.join(refs)
                        writer.writerow([main_concept, term, refs_str])
    
    def _export_html(self, hierarchical_index: Dict[str, any], 
                     output_path: str, format_type: str):
        """Export index as HTML with styling."""
        # Check if 3-level or 2-level
        first_entries = next(iter(hierarchical_index.values()))
        is_3level = first_entries and isinstance(next(iter(first_entries.values())), dict)
        
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Hierarchical Index</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        .container {{
            background: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            margin-bottom: 30px;
            font-size: 2em;
        }}
        .category {{
            margin-bottom: 30px;
            padding: 20px;
            background: #f8f9fa;
            border-radius: 6px;
            border-left: 4px solid #3498db;
        }}
        .category-title {{
            font-size: 1.5em;
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 15px;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        .subcategory {{
            margin: 15px 0 15px 20px;
            padding: 12px;
            background: white;
            border-radius: 4px;
            border-left: 3px solid #95a5a6;
        }}
        .subcategory-title {{
            font-weight: 600;
            color: #34495e;
            margin-bottom: 8px;
            font-size: 1.1em;
        }}
        .entry {{
            margin: 8px 0 8px 20px;
            padding: 8px 12px;
            background: #fff;
            border-radius: 3px;
            border-left: 2px solid #bdc3c7;
        }}
        .term {{
            font-weight: 500;
            color: #2c3e50;
            display: inline-block;
            min-width: 200px;
        }}
        .references {{
            color: #7f8c8d;
            font-family: 'Courier New', monospace;
            font-size: 0.95em;
        }}
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #e0e0e0;
            text-align: center;
            color: #7f8c8d;
            font-size: 0.9em;
        }}
        @media print {{
            body {{
                background: white;
            }}
            .container {{
                box-shadow: none;
                padding: 20px;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Hierarchical Index</h1>
"""
        
        # Generate HTML content
        for main_concept in sorted(hierarchical_index.keys()):
            html_content += f'        <div class="category">\n'
            html_content += f'            <div class="category-title">{main_concept.upper()}</div>\n'
            
            entries = hierarchical_index[main_concept]
            
            if is_3level:
                # 3-level hierarchy
                for sub_cluster in sorted(entries.keys()):
                    html_content += f'            <div class="subcategory">\n'
                    html_content += f'                <div class="subcategory-title">[{sub_cluster}]</div>\n'
                    
                    sub_entries = entries[sub_cluster]
                    for term in sorted(sub_entries.keys()):
                        refs = sub_entries[term]
                        refs_str = '; '.join(refs)
                        html_content += f'                <div class="entry">\n'
                        html_content += f'                    <span class="term">{term}</span>\n'
                        html_content += f'                    <span class="references">{refs_str}</span>\n'
                        html_content += f'                </div>\n'
                    
                    html_content += f'            </div>\n'
            else:
                # 2-level hierarchy
                for term in sorted(entries.keys()):
                    refs = entries[term]
                    refs_str = '; '.join(refs)
                    html_content += f'            <div class="entry">\n'
                    html_content += f'                <span class="term">{term}</span>\n'
                    html_content += f'                <span class="references">{refs_str}</span>\n'
                    html_content += f'            </div>\n'
            
            html_content += f'        </div>\n'
        
        html_content += """        <div class="footer">
            Generated by LexiFinder CLI
        </div>
    </div>
</body>
</html>
"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

    def mark_document(self, filepath: str, hierarchical_nouns: any, output_path: str):
        """
        Mark keywords in DOCX/ODT document with index entry fields.
        
        Args:
            filepath: Path to input document
            hierarchical_nouns: Dictionary mapping keywords to their correlated nouns (2 or 3 levels)
            output_path: Path to save marked document
        """
        # Flatten hierarchy for marking (we only mark the actual nouns)
        flat_hierarchy = self._flatten_hierarchy(hierarchical_nouns)
        
        file_type = self.detect_file_type(filepath)
        
        if file_type == 'docx':
            self._mark_docx(filepath, flat_hierarchy, output_path)
        elif file_type == 'odt':
            self._mark_odt(filepath, flat_hierarchy, output_path)
        else:
            raise ValueError("Document marking is only supported for DOCX and ODT files.")
    
    def _mark_docx(self, filepath: str, hierarchical_nouns: Dict[str, List[str]], output_path: str):
        """Mark keywords in DOCX document with XE (Index Entry) fields."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        print(f"Marking keywords in DOCX document...")
        
        doc = DocxDocument(filepath)
        
        # Build a flat list of all nouns with their main keyword
        noun_to_keywords = {}  # noun -> [list of keywords it belongs to]
        for keyword, nouns in hierarchical_nouns.items():
            for noun in nouns:
                if noun.lower() not in noun_to_keywords:
                    noun_to_keywords[noun.lower()] = []
                noun_to_keywords[noun.lower()].append(keyword)
        
        marked_count = 0
        
        # Process each paragraph
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.lower()
            
            # Check if any noun appears in this paragraph
            for noun_lower, keywords in noun_to_keywords.items():
                if noun_lower in para_text:
                    # Mark each occurrence of this noun
                    for keyword in keywords:
                        # Add XE field at the end of the paragraph
                        # Format: MainEntry:SubEntry
                        self._insert_xe_field(paragraph, keyword, noun_lower.capitalize())
                        marked_count += 1
            
            if (para_idx + 1) % 50 == 0:
                print(f"  Processed {para_idx + 1}/{len(doc.paragraphs)} paragraphs...")
        
        # Save marked document
        doc.save(output_path)
        print(f"  Document marked successfully: {marked_count} index entries added.")
        print(f"  Saved to: {output_path}")
        print(f"\n  To generate the index in Word:")
        print(f"    1. Open the marked document")
        print(f"    2. Place cursor where you want the index")
        print(f"    3. Go to: References → Insert Index → OK")
    
    def _insert_xe_field(self, paragraph, main_entry: str, sub_entry: str):
        """
        Insert an XE (Index Entry) field into a paragraph.
        
        Format: { XE "MainEntry:SubEntry" }
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Create a new run at the end of the paragraph
        run = paragraph.add_run()
        
        # Create fldChar element for field start
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        
        # Create instrText element with XE field code
        # Format: XE "MainEntry:SubEntry"
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f' XE "{main_entry}:{sub_entry}" '
        
        # Create fldChar element for field end
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        
        # Add elements to run
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)
    
    def _mark_odt(self, filepath: str, hierarchical_nouns: Dict[str, List[str]], output_path: str):
        """Mark keywords in ODT document with alphabetical index fields."""
        from odf.opendocument import load as load_odt
        from odf import text as odf_text
        from odf.text import AlphabeticalIndexMarkStart, AlphabeticalIndexMarkEnd
        import re
        
        print(f"Marking keywords in ODT document...")
        
        doc = load_odt(filepath)
        
        # Build noun to keywords mapping
        noun_to_keywords = {}
        for keyword, nouns in hierarchical_nouns.items():
            for noun in nouns:
                if noun.lower() not in noun_to_keywords:
                    noun_to_keywords[noun.lower()] = []
                noun_to_keywords[noun.lower()].append(keyword)
        
        marked_count = 0
        
        # Process paragraphs
        paragraphs = doc.getElementsByType(odf_text.P)
        for para_idx, paragraph in enumerate(paragraphs):
            para_text = str(paragraph)
            para_text_clean = re.sub('<[^>]+>', '', para_text).strip().lower()
            
            # Check for nouns in this paragraph
            for noun_lower, keywords in noun_to_keywords.items():
                if noun_lower in para_text_clean:
                    # Add index mark for each keyword
                    for keyword in keywords:
                        # Create alphabetical index mark
                        mark_start = AlphabeticalIndexMarkStart(
                            key1=keyword,
                            key2=noun_lower.capitalize()
                        )
                        mark_end = AlphabeticalIndexMarkEnd(id=f"idx_{marked_count}")
                        
                        # Add marks to paragraph
                        paragraph.addElement(mark_start)
                        paragraph.addElement(mark_end)
                        marked_count += 1
            
            if (para_idx + 1) % 50 == 0:
                print(f"  Processed {para_idx + 1}/{len(paragraphs)} paragraphs...")
        
        # Save marked document
        doc.save(output_path)
        print(f"  Document marked successfully: {marked_count} index entries added.")
        print(f"  Saved to: {output_path}")
        print(f"\n  To generate the index in LibreOffice:")
        print(f"    1. Open the marked document")
        print(f"    2. Place cursor where you want the index")
        print(f"    3. Go to: Insert → Table of Contents and Index → Alphabetical Index → OK")

    def process(self, filepath: str, output_path: str, keywords: List[str] = None, 
                mode: str = 'page', mark_document: bool = False, 
                strategy: str = 'keywords', export_format: str = 'txt', 
                **strategy_options):
        """
        Main processing pipeline for hierarchical index generation.
        
        Args:
            filepath: Path to input document
            output_path: Path to output index file
            keywords: List of keywords (required for keywords/hybrid strategies)
            mode: 'page' or 'paragraph'
            mark_document: If True, mark DOCX/ODT with index fields instead of creating text index
            strategy: Indexing strategy ('keywords', 'auto', 'hybrid', 'frequent')
            export_format: Output format ('txt', 'json', 'csv', 'html')
            **strategy_options: Additional options (max_per_category, clusters, subclusters, top)
        """
        try:
            file_type = self.detect_file_type(filepath)
            
            # Validate keywords for strategies that need them
            if strategy in ['keywords', 'hybrid'] and not keywords:
                print(f"✗ Error: Strategy '{strategy}' requires keywords (-k parameter).")
                sys.exit(1)
            
            # Check if marking is requested but file type doesn't support it
            if mark_document and file_type == 'pdf':
                print("⚠ Warning: Document marking is only supported for DOCX and ODT files.")
                print("  Proceeding with text index generation instead.\n")
                mark_document = False
            
            # Step 1: Extract all nouns from document
            nouns = self.extract_nouns(filepath)
            
            # Get full text for strategies that need it
            text = ""
            if strategy == 'frequent':
                if file_type == 'pdf':
                    text = self.reader.read_pdf_text(filepath)
                elif file_type == 'docx':
                    text = self.reader.read_docx_text(filepath)
                elif file_type == 'odt':
                    text = self.reader.read_odt_text(filepath)
            
            # Step 2: Build hierarchical structure using selected strategy
            hierarchical_nouns = self.find_correlated_nouns(
                nouns, 
                keywords=keywords, 
                text=text, 
                strategy=strategy, 
                **strategy_options
            )
            
            # Check if any concept has sub-entries
            # Handle both 2-level and 3-level hierarchies
            if hierarchical_nouns and isinstance(next(iter(hierarchical_nouns.values())), dict) and \
               strategy == 'hybrid':
                # 3-level: count all nouns in all sub-clusters
                total_subentries = sum(
                    sum(len(v) for v in cat.values()) 
                    for cat in hierarchical_nouns.values()
                )
            else:
                # 2-level: count nouns in each category
                total_subentries = sum(len(entries) for entries in hierarchical_nouns.values())
            if total_subentries == 0:
                print("\n⚠ Warning: No nouns matched the similarity criteria for any concept.")
                print("Try lowering the similarity threshold or using different keywords.")
                return
            
            # If marking is requested for DOCX/ODT
            if mark_document and file_type in ['docx', 'odt']:
                # Generate marked document filename
                base_name = os.path.splitext(filepath)[0]
                ext = os.path.splitext(filepath)[1]
                marked_path = f"{base_name}_marked{ext}"
                
                print(f"\n{'='*70}")
                print(f"Creating marked document for index generation...")
                print(f"{'='*70}\n")
                
                self.mark_document(filepath, hierarchical_nouns, marked_path)
                
                # Print summary
                print(f"\n✓ Document marking completed successfully!")
                print(f"  Document type: {file_type.upper()}")
                print(f"  Main concepts: {len(hierarchical_nouns)}")
                print(f"  Total sub-entries: {total_subentries}")
                
                # Print distribution
                for keyword in sorted(hierarchical_nouns.keys()):
                    count = len(hierarchical_nouns[keyword])
                    if count > 0:
                        print(f"    • {keyword}: {count} sub-entries")
                
                return
            
            # Otherwise, proceed with text index generation
            # Step 3: Build location index for all sub-entries
            hierarchical_index, format_type = self.extract_occurrences(
                filepath, hierarchical_nouns, mode
            )
            
            # Step 4: Write hierarchical index to file
            self.write_index(hierarchical_index, output_path, format_type, export_format)
            
            # Print summary
            if format_type == 'page':
                index_mode = "pages"
            elif format_type == 'page_para':
                index_mode = "pages and paragraphs"
            else:  # chapter_para
                index_mode = "chapters and paragraphs"
            
            print(f"\n✓ Process completed successfully!")
            print(f"  Document type: {file_type.upper()}")
            print(f"  Indexing mode: {index_mode}")
            print(f"  Main concepts: {len(hierarchical_index)}")
            print(f"  Total sub-entries: {total_subentries}")
            
            # Print distribution
            for keyword in sorted(hierarchical_index.keys()):
                count = len(hierarchical_index[keyword])
                if count > 0:
                    print(f"    • {keyword}: {count} sub-entries")
            
        except Exception as e:
            print(f"\n✗ Error: {e}", file=sys.stderr)
            import traceback
            traceback.print_exc()
            sys.exit(1)
    
    def process_batch(self, input_dir: str, output_dir: str, 
                     pattern: str = "*", keywords: List[str] = None,
                     mode: str = 'page', mark_document: bool = False,
                     strategy: str = 'keywords', export_format: str = 'txt',
                     **strategy_options):
        """
        Batch process multiple documents with same configuration.
        
        Args:
            input_dir: Directory containing input documents
            output_dir: Directory for output files
            pattern: File pattern (default: "*" for all supported files)
            keywords: List of keywords (required for keywords/hybrid strategies)
            mode: 'page' or 'paragraph'
            mark_document: If True, mark DOCX/ODT with index fields
            strategy: Indexing strategy
            export_format: Output format ('txt', 'json', 'csv', 'html')
            **strategy_options: Additional options
        """
        print("="*70)
        print("BATCH PROCESSING MODE")
        print("="*70)
        print(f"Input directory:  {input_dir}")
        print(f"Output directory: {output_dir}")
        print(f"Pattern:          {pattern}")
        print(f"Strategy:         {strategy}")
        if keywords:
            print(f"Keywords:         {', '.join(keywords)}")
        print(f"Export format:    {export_format}")
        print("="*70)
        print()
        if mode == 'page':
            print("ℹ  Note: for DOCX/ODT files in this batch, paragraph mode will be used automatically.")
            print()
        # Ensure output directory exists
        if not os.path.exists(output_dir):
            print(f"⚠ Output directory does not exist: {output_dir}")
            try:
                answer = input("Do you want to create it? [y/N] ").strip().lower()
            except (EOFError, KeyboardInterrupt):
                answer = 'n'
            if answer in ('y', 'yes'):
                os.makedirs(output_dir, exist_ok=True)
                print(f"✓ Created output directory: {output_dir}\n")
            else:
                print("Aborted. Output directory was not created.")
                return False
        
        # Find all matching files
        supported_extensions = ['.pdf', '.docx', '.odt']
        all_files = []
        
        # If pattern contains extension, use it directly
        if any(pattern.endswith(ext) for ext in supported_extensions):
            search_pattern = os.path.join(input_dir, pattern)
            all_files = glob.glob(search_pattern)
        else:
            # Search for all supported file types
            for ext in supported_extensions:
                search_pattern = os.path.join(input_dir, f"{pattern}{ext}")
                all_files.extend(glob.glob(search_pattern))
        
        if not all_files:
            print(f"✗ No files found matching pattern '{pattern}' in '{input_dir}'")
            print(f"  Supported formats: {', '.join(supported_extensions)}")
            sys.exit(1)
        
        print(f"Found {len(all_files)} file(s) to process:\n")
        for f in all_files:
            print(f"  • {os.path.basename(f)}")
        print()
        
        # Process each file
        results = {
            'success': [],
            'failed': [],
            'skipped': []
        }
        
        for idx, input_file in enumerate(all_files, 1):
            filename = os.path.basename(input_file)
            name_without_ext = os.path.splitext(filename)[0]
            
            # Determine output filename
            output_file = os.path.join(output_dir, f"{name_without_ext}-index.{export_format}")
            
            print(f"\n[{idx}/{len(all_files)}] Processing: {filename}")
            print("-"*70)
            
            try:
                # Determine effective mode for this file
                file_ext = os.path.splitext(input_file)[1].lower()
                effective_mode = mode
                if file_ext in ['.docx', '.odt'] and mode == 'page':
                    effective_mode = 'paragraph'

                # Process the file
                self.process(
                    input_file,
                    output_file,
                    keywords=keywords,
                    mode=effective_mode,
                    mark_document=mark_document,
                    strategy=strategy,
                    export_format=export_format,
                    **strategy_options
                )
                results['success'].append(filename)
                print(f"✓ Completed: {filename} → {os.path.basename(output_file)}")
                
            except Exception as e:
                results['failed'].append((filename, str(e)))
                print(f"✗ Failed: {filename}")
                print(f"  Error: {e}")
                # Continue with next file
        
        # Print summary
        print("\n" + "="*70)
        print("BATCH PROCESSING SUMMARY")
        print("="*70)
        print(f"Total files:      {len(all_files)}")
        print(f"✓ Successful:     {len(results['success'])}")
        print(f"✗ Failed:         {len(results['failed'])}")
        if results['skipped']:
            print(f"⊘ Skipped:        {len(results['skipped'])}")
        print("="*70)
        
        if results['success']:
            print("\n✓ Successfully processed files:")
            for f in results['success']:
                print(f"  • {f}")
        
        if results['failed']:
            print("\n✗ Failed files:")
            for f, error in results['failed']:
                print(f"  • {f}: {error}")
        
        print(f"\n✓ Output files saved to: {output_dir}")
        
        # Return success if at least one file processed
        return len(results['success']) > 0


def parse_keywords(keywords_str: str) -> List[str]:
    """
    Parse keywords from semicolon-separated string.
    
    Args:
        keywords_str: String containing keywords separated by semicolons
        
    Returns:
        List of cleaned keywords
    """
    keywords = [kw.strip() for kw in keywords_str.split(';') if kw.strip()]
    if not keywords:
        raise ValueError("No valid keywords provided.")
    return keywords


def main():
    parser = argparse.ArgumentParser(
        description="LexiFinder CLI - Generate a hierarchical analytic index from PDF, DOCX, or ODT.\n"
                    "Creates a two-level index: main concepts (your keywords) and related sub-entries.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # STRATEGY 1: Keywords (default - user-defined categories)
  %(prog)s -i document.pdf -o index.txt -k "science; technology" -m page
  %(prog)s -i thesis.docx -o index.txt -k "AI; machine learning; neural networks"
  
  # STRATEGY 2: Hybrid (keywords + automatic sub-clustering, 3 levels)
  %(prog)s -i document.pdf -o index.txt -k "AI; ML" --strategy hybrid --subclusters 3
  
  # STRATEGY 3: Auto (fully automatic clustering)
  %(prog)s -i document.pdf -o index.txt --strategy auto --clusters 10
  
  # STRATEGY 4: Frequent (most frequent terms as categories)
  %(prog)s -i document.pdf -o index.txt --strategy frequent --top 15
  
  # Advanced options
  %(prog)s -i doc.pdf -o index.txt -k "AI" --strategy hybrid --max-per-category 20
  
  # Smart filtering: exclude generic words
  %(prog)s -i doc.pdf -o index.txt -k "AI" --exclude-generic
  
  # Smart filtering: only terms with 3+ occurrences
  %(prog)s -i doc.pdf -o index.txt -k "AI" --min-occurrences 3
  
  # Combine filters for clean index
  %(prog)s -i doc.pdf -o index.txt -k "AI" --exclude-generic --min-occurrences 5
  
  # Mark DOCX for native Word index generation (works with any strategy)
  %(prog)s -i thesis.docx -o temp.txt -k "science" --strategy keywords -x
  
  # Use specific model
  %(prog)s -i documento.pdf -o indice.txt -k "parole" --model it_core_news_md

Model management:
  %(prog)s --list-models              # List available models
  %(prog)s --list-installed           # List installed models
  %(prog)s --download-model it_core_news_md   # Download a model
  %(prog)s --delete-model it_core_news_md     # Delete a model

Configuration management:
  # Save configuration for reuse
  %(prog)s -i doc.pdf -o index.txt -k "AI; ML" --strategy hybrid --save-config my-config.json
  
  # Load saved configuration
  %(prog)s -i chapter2.pdf -o ch2-index.txt --load-config my-config.json
  
  # Load config and override specific parameters
  %(prog)s -i new-doc.pdf -o new-index.txt --load-config my-config.json -t 0.7

Export formats:
  # Plain text (default)
  %(prog)s -i doc.pdf -o index.txt -k "AI" -f txt
  
  # JSON (structured data)
  %(prog)s -i doc.pdf -o index.json -k "AI" -f json
  
  # CSV (spreadsheet-friendly)
  %(prog)s -i doc.pdf -o index.csv -k "AI" -f csv
  
  # HTML (web-ready with styling)
  %(prog)s -i doc.pdf -o index.html -k "AI" -f html

Batch processing:
  # Process all PDFs in a directory
  %(prog)s --batch-dir ./documents --output-dir ./indexes -k "AI; ML"
  
  # Process specific pattern
  %(prog)s --batch-dir ./chapters --output-dir ./indexes --pattern "chapter*.pdf" -k "topics"
  
  # Batch with config file
  %(prog)s --batch-dir ./docs --output-dir ./output --load-config config.json
  
  # Batch with all formats
  %(prog)s --batch-dir ./docs --output-dir ./html-indexes -k "AI" -f html

Strategy Details:
  keywords  - 2 levels: Your keywords → Similar nouns
  hybrid    - 3 levels: Your keywords → Auto sub-groups → Similar nouns  
  auto      - 2 levels: Auto categories → Similar nouns
  frequent  - 2 levels: Frequent terms → Similar nouns

Output format examples:
  Keywords/Auto/Frequent (2 levels):
    MAIN CATEGORY
        sub-entry: p.5; p.18; p.34
        
  Hybrid (3 levels):
    MAIN KEYWORD
      [Sub-group 1]
        noun: p.5; p.18
      [Sub-group 2]
        noun: p.10; p.25

Supported formats: .pdf, .docx, .odt
        """
    )
    
    parser.add_argument(
        '-i', '--input',
        required=False,
        metavar='FILE',
        help='Path to input file (.pdf, .docx, or .odt)'
    )
    
    parser.add_argument(
        '-o', '--output',
        required=False,
        metavar='TXT',
        help='Path to output index file (.txt)'
    )
    
    # Batch processing parameters
    parser.add_argument(
        '--batch-dir',
        metavar='DIR',
        help='Process all documents in directory (batch mode)'
    )
    
    parser.add_argument(
        '--output-dir',
        metavar='DIR',
        help='Output directory for batch processing (required with --batch-dir)'
    )
    
    parser.add_argument(
        '--pattern',
        default='*',
        metavar='PATTERN',
        help='File pattern for batch processing (default: "*" for all supported files)'
    )
    
    parser.add_argument(
        '-k', '--keywords',
        required=False,
        metavar='KEYWORDS',
        help='Keywords separated by semicolons (required for keywords/hybrid strategies)'
    )
    
    parser.add_argument(
        '-t', '--threshold',
        type=float,
        default=0.5,
        metavar='FLOAT',
        help='Similarity threshold (0.0 to 1.0, default: 0.5)'
    )
    
    parser.add_argument(
        '-m', '--mode',
        choices=['page', 'paragraph'],
        default='page',
        help='Indexing mode: "page" or "paragraph" (default: page for PDF, paragraph for DOCX/ODT)'
    )
    
    parser.add_argument(
        '-x', '--mark',
        action='store_true',
        help='Mark DOCX/ODT document with index fields for native index generation (DOCX/ODT only)'
    )
    
    parser.add_argument(
        '-f', '--export-format',
        choices=['txt', 'json', 'csv', 'html'],
        default='txt',
        metavar='FORMAT',
        help='Output format: txt (default), json, csv, or html'
    )
    
    # Config management
    parser.add_argument(
        '--save-config',
        metavar='FILE',
        help='Save current configuration to a JSON file'
    )
    
    parser.add_argument(
        '--load-config',
        metavar='FILE',
        help='Load configuration from a JSON file'
    )
    
    # Strategy parameters
    parser.add_argument(
        '--strategy',
        choices=['keywords', 'auto', 'hybrid', 'frequent'],
        default='keywords',
        help='Indexing strategy (default: keywords). See --help for details on each strategy.'
    )
    
    parser.add_argument(
        '--max-per-category',
        type=int,
        default=30,
        metavar='N',
        help='Maximum sub-entries per category (default: 30)'
    )
    
    parser.add_argument(
        '--clusters',
        type=int,
        default=8,
        metavar='N',
        help='Number of clusters for auto strategy (default: 8)'
    )
    
    parser.add_argument(
        '--subclusters',
        type=int,
        default=3,
        metavar='N',
        help='Number of sub-clusters for hybrid strategy (default: 3)'
    )
    
    parser.add_argument(
        '--top',
        type=int,
        default=15,
        metavar='N',
        help='Number of top frequent terms for frequent strategy (default: 15)'
    )
    
    # Smart filtering parameters
    parser.add_argument(
        '--exclude-generic',
        action='store_true',
        help='Exclude generic/common words (thing, way, part, aspect, etc.)'
    )
    
    parser.add_argument(
        '--min-occurrences',
        type=int,
        default=1,
        metavar='N',
        help='Minimum occurrences required to include a term (default: 1)'
    )
    
    parser.add_argument(
        '--model',
        default=None,
        metavar='MODEL',
        help=f'SpaCy model to use (default: {LexiFinderCLI.DEFAULT_MODEL}). Use --list-models to see available options.'
    )
    
    parser.add_argument(
        '--list-models',
        action='store_true',
        help='List all available spaCy models and exit'
    )
    
    parser.add_argument(
        '--download-model',
        metavar='MODEL',
        help='Download a spaCy model and exit'
    )
    
    parser.add_argument(
        '--delete-model',
        metavar='MODEL',
        help='Delete an installed spaCy model and exit'
    )
    
    parser.add_argument(
        '--list-installed',
        action='store_true',
        help='List all installed spaCy models and exit'
    )
    
    parser.add_argument(
        '--version',
        action='version',
        version=(
            f'LexiFinder v{APP_VERSION}\n'
            f'by {APP_AUTHOR}\n'
            f'Repository : {APP_REPO}\n'
            f'Webpage    : {APP_WEB}\n'
            f'Donate     : {APP_PAYPAL}'
        )
    )

    parser.add_argument(
        '--gui',
        action='store_true',
        help='Launch the graphical user interface'
    )

    args = parser.parse_args()

    # Handle --gui
    if args.gui:
        run_gui()
        sys.exit(0)
    
    # Handle --list-models
    if args.list_models:
        LexiFinderCLI.list_available_models()
        sys.exit(0)
    
    # Handle --list-installed
    if args.list_installed:
        LexiFinderCLI.list_installed_models()
        sys.exit(0)
    
    # Handle --download-model
    if args.download_model:
        print(f"Downloading model: {args.download_model}")
        success = LexiFinderCLI.download_model(args.download_model)
        sys.exit(0 if success else 1)
    
    # Handle --delete-model
    if args.delete_model:
        print(f"Deleting model: {args.delete_model}")
        success = LexiFinderCLI.delete_model(args.delete_model)
        sys.exit(0 if success else 1)
    
    # Load configuration file if specified (overrides defaults)
    if args.load_config:
        config = LexiFinderCLI.load_config(args.load_config)
        
        # Apply config values (command-line args override config)
        if 'keywords' in config and not args.keywords:
            args.keywords = config['keywords']
        if 'threshold' in config and args.threshold == 0.5:  # default value
            args.threshold = config['threshold']
        if 'mode' in config and args.mode == 'page':  # default value
            args.mode = config['mode']
        if 'strategy' in config and args.strategy == 'keywords':  # default value
            args.strategy = config['strategy']
        if 'max_per_category' in config and args.max_per_category == 30:  # default value
            args.max_per_category = config['max_per_category']
        if 'clusters' in config and args.clusters == 8:  # default value
            args.clusters = config['clusters']
        if 'subclusters' in config and args.subclusters == 3:  # default value
            args.subclusters = config['subclusters']
        if 'top' in config and args.top == 15:  # default value
            args.top = config['top']
        if 'exclude_generic' in config and not args.exclude_generic:
            args.exclude_generic = config['exclude_generic']
        if 'min_occurrences' in config and args.min_occurrences == 1:  # default value
            args.min_occurrences = config['min_occurrences']
        if 'model' in config and not args.model:
            args.model = config['model']
        if 'export_format' in config and args.export_format == 'txt':  # default value
            args.export_format = config['export_format']
        if 'mark' in config and not args.mark:
            args.mark = config['mark']
    
    # Determine operation mode: batch or single file
    batch_mode = args.batch_dir is not None
    
    if batch_mode:
        # Batch mode validation
        if not args.output_dir:
            parser.error("--output-dir is required when using --batch-dir")
        if not os.path.isdir(args.batch_dir):
            parser.error(f"Batch directory not found: {args.batch_dir}")
    else:
        # Single file mode validation
        if not args.input:
            parser.error("the following arguments are required: -i/--input (or use --batch-dir for batch mode)")
        if not args.output:
            parser.error("the following arguments are required: -o/--output (or use --output-dir for batch mode)")
        
        # DOCX/ODT do not have pages: force paragraph mode before printing the banner
        ext = os.path.splitext(args.input)[1].lower()
        if ext in ['.docx', '.odt']:
            args.mode = 'paragraph'
    
    # Validate threshold
    if not 0.0 <= args.threshold <= 1.0:
        parser.error("Threshold must be between 0.0 and 1.0")
    
    # Parse keywords (if provided)
    keywords = None
    if args.keywords:
        try:
            keywords = parse_keywords(args.keywords)
        except ValueError as e:
            parser.error(str(e))
    
    # Validate keywords for strategies that need them
    if args.strategy in ['keywords', 'hybrid'] and not keywords:
        parser.error(f"Strategy '{args.strategy}' requires keywords (-k parameter)")
    
    # Execute based on mode
    if batch_mode:
        # Batch processing
        finder = LexiFinderCLI(similarity_threshold=args.threshold, model_name=args.model)
        success = finder.process_batch(
            args.batch_dir,
            args.output_dir,
            pattern=args.pattern,
            keywords=keywords,
            mode=args.mode,
            mark_document=args.mark,
            strategy=args.strategy,
            export_format=args.export_format,
            max_per_category=args.max_per_category,
            clusters=args.clusters,
            subclusters=args.subclusters,
            top=args.top,
            exclude_generic=args.exclude_generic,
            min_occurrences=args.min_occurrences
        )
        
        # Save configuration if requested
        if args.save_config:
            LexiFinderCLI.save_config(args.save_config, args)
        
        sys.exit(0 if success else 1)
    
    # Single file mode (existing code)
    # Ensure output has correct extension
    if not args.output.lower().endswith(f'.{args.export_format}'):
        args.output = os.path.splitext(args.output)[0] + f'.{args.export_format}'
    
    # Print configuration
    print("=" * 70)
    print("LexiFinder CLI - Hierarchical Analytic Index Generator")
    print("=" * 70)
    print(f"Input file:    {args.input}")
    print(f"Output file:   {args.output}")
    print(f"Strategy:      {args.strategy}")
    if keywords:
        print(f"Keywords:      {', '.join(keywords)}")
    print(f"Threshold:     {args.threshold}")
    print(f"Mode:          {args.mode}")
    if args.model:
        print(f"Model:         {args.model}")
    if args.strategy == 'auto':
        print(f"Clusters:      {args.clusters}")
    elif args.strategy == 'hybrid':
        print(f"Sub-clusters:  {args.subclusters}")
    elif args.strategy == 'frequent':
        print(f"Top terms:     {args.top}")
    print(f"Max/category:  {args.max_per_category}")
    if args.exclude_generic:
        print(f"Smart filter:  Exclude generic words")
    if args.min_occurrences > 1:
        print(f"Min occurs:    {args.min_occurrences}")
    if args.mark:
        print(f"Mark document: YES")
    print("=" * 70)
    print()
    
    # Run the process
    finder = LexiFinderCLI(similarity_threshold=args.threshold, model_name=args.model)
    finder.process(
        args.input, 
        args.output, 
        keywords=keywords, 
        mode=args.mode, 
        mark_document=args.mark,
        strategy=args.strategy,
        export_format=args.export_format,
        max_per_category=args.max_per_category,
        clusters=args.clusters,
        subclusters=args.subclusters,
        top=args.top,
        exclude_generic=args.exclude_generic,
        min_occurrences=args.min_occurrences
    )
    
    # Save configuration if requested
    if args.save_config:
        LexiFinderCLI.save_config(args.save_config, args)


def run_gui():
    """Launch the LexiFinder graphical user interface using Flet."""
    try:
        import flet as ft
    except ImportError:
        print("✗ Flet is not installed. Install it with: pip install flet")
        sys.exit(1)

    import threading
    import io
    import json
    import time as _time

    # ─── stdout/stderr redirector ─────────────────────────────────────────────
    import queue as _queue
    import time as _time

    class GuiLogger(io.TextIOBase):
        """Redirect stdout/stderr to a Flet ListView, filtering verbose noise.

        UI updates are batched in a dedicated flush loop that fires every
        UI_REFRESH_INTERVAL seconds.  The worker thread therefore never blocks
        on page.update(), keeping both the backend and the interface responsive
        even on slow machines.
        """

        UI_REFRESH_INTERVAL = 0.10   # seconds between UI refreshes (≈10 fps)

        # Lines matching these patterns are suppressed in the log
        _NOISE_PATTERNS = [
            # tqdm progress bars — identified by the "|  N%|" pattern.
            # _on_step reads the percentage from these lines to update the
            # progress bar, but we don't want them cluttering the log.
            lambda l: ("|" in l and "%" in l),
            lambda l: ("it/s" in l or "s/it" in l),
            # Long separator lines (====== or ------)
            lambda l: len(l) >= 20 and all(c in ("=", "-") for c in l.strip()) and l.strip() != "",
            # tqdm block-fill bar characters
            lambda l: "█" in l or "▏" in l or "▎" in l or "▍" in l,
        ]

        def __init__(self, log_view, page, on_step=None):
            self._log   = log_view
            self._page  = page
            self._buf   = ""
            self._on_step = on_step

            # Thread-safe queue: worker pushes ft.Text controls,
            # flush loop drains it and calls page.update() in bulk.
            self._item_queue: _queue.Queue = _queue.Queue()

            # Flag: True when progress-bar values have changed but
            # page.update() has not been called yet.
            self._ui_dirty = threading.Event()

            # Stop signal for the flush loop
            self._stop_event = threading.Event()

            # Start the dedicated UI-update thread
            self._flush_thread = threading.Thread(
                target=self._flush_loop, daemon=True, name="GuiLogger-flush"
            )
            self._flush_thread.start()

        # ── public API ────────────────────────────────────────────────────────

        def set_step_callback(self, cb):
            self._on_step = cb

        def schedule_update(self):
            """Signal the flush loop that a UI refresh is needed (e.g. from
            progress-bar updates).  Non-blocking: safe to call from any thread."""
            self._ui_dirty.set()

        def stop(self):
            """Gracefully stop the flush loop (called when the app exits)."""
            self._stop_event.set()

        # ── io.TextIOBase interface ───────────────────────────────────────────

        def write(self, text: str) -> int:
            if not text:
                return 0
            # Normalise \r\n → \n first, then standalone \r → \n.
            # This is the key fix for tqdm compatibility: tqdm writes its
            # in-place progress updates terminated by \r (not \n), so without
            # this step the buffer never sees a complete line, _append() is
            # never called during long loops, and the UI freezes until the
            # very end when a final \n is eventually written.
            text = text.replace("\r\n", "\n").replace("\r", "\n")
            self._buf += text
            while "\n" in self._buf:
                line, self._buf = self._buf.split("\n", 1)
                self._append(line)
            return len(text)

        def flush(self):
            if self._buf:
                self._append(self._buf)
                self._buf = ""

        # ── internal helpers ─────────────────────────────────────────────────

        def _should_suppress(self, line: str) -> bool:
            stripped = line.strip()
            if not stripped:
                return True
            for check in self._NOISE_PATTERNS:
                if check(stripped):
                    return True
            return False

        def _append(self, line: str):
            # Notify progress callback regardless of suppression.
            # _on_step should only *set values* on controls; the flush loop
            # will call page.update() for us.
            if self._on_step:
                try:
                    self._on_step(line)
                except Exception:
                    pass

            if self._should_suppress(line):
                return

            try:
                is_dark = (self._page.platform_brightness == ft.Brightness.DARK
                           or self._page.theme_mode == ft.ThemeMode.DARK)
            except Exception:
                is_dark = True
            if is_dark:
                _c_default, _c_ok, _c_err, _c_warn, _c_info = (
                    "#DDDDDD", "#66BB6A", "#EF5350", "#FFA726", "#90CAF9")
            else:
                _c_default, _c_ok, _c_err, _c_warn, _c_info = (
                    "#222222", "#2E7D32", "#C62828", "#E65100", "#1565C0")
            color = _c_default
            stripped = line.strip()
            if stripped.startswith("✓") or "successfully" in stripped.lower():
                color = _c_ok
            elif stripped.startswith("✗") or "error" in stripped.lower():
                color = _c_err
            elif stripped.startswith("⚠") or "warning" in stripped.lower():
                color = _c_warn
            elif stripped.startswith("ℹ"):
                color = _c_info

            # Enqueue the control; do NOT touch the UI from this thread
            self._item_queue.put(
                ft.Text(line, size=12, color=color,
                        font_family="monospace", selectable=True)
            )

        def _flush_loop(self):
            """Background thread: drain the log queue and refresh the UI in bulk.
            
            Progress-bar / label updates are pushed immediately via run_task()
            in _on_step; this loop only handles log-item batching.
            """
            while not self._stop_event.is_set():
                # Wait for the dirty flag OR the periodic interval, whichever
                # comes first.  This makes the loop react quickly to bursts of
                # log output without hammering the event loop every 100 ms when
                # there is nothing to do.
                self._ui_dirty.wait(timeout=self.UI_REFRESH_INTERVAL)

                # Drain all pending log items
                items = []
                try:
                    while True:
                        items.append(self._item_queue.get_nowait())
                except _queue.Empty:
                    pass

                has_log    = bool(items)
                has_dirty  = self._ui_dirty.is_set()

                if items:
                    self._log.controls.extend(items)

                if has_dirty:
                    self._ui_dirty.clear()

                if has_log or has_dirty:
                    # Schedule the refresh on Flet's own event loop so it
                    # executes in the same thread that drives the WebSocket —
                    # this avoids the "update called from a non-Flet thread"
                    # latency / silently-dropped-update problem.
                    _log_ref  = self._log
                    _page_ref = self._page
                    async def _do_log_update():
                        try:
                            await _page_ref.update_async()
                        except Exception:
                            try:
                                _page_ref.update()
                            except Exception:
                                pass
                    try:
                        self._page.run_task(_do_log_update)
                    except Exception:
                        # Older Flet build without run_task – fall back to the
                        # direct (possibly-slow) synchronous call.
                        try:
                            self._page.update()
                        except Exception:
                            pass

    # ─── Flet main ────────────────────────────────────────────────────────────
    def main_gui(page: ft.Page):
        page.title = "LexiFinder – Hierarchical Index Generator"
        page.theme_mode = ft.ThemeMode.SYSTEM
        page.window.width = 1100
        page.window.height = 780
        page.window.min_width = 800
        page.window.min_height = 600
        page.padding = 0

        # ── state ──────────────────────────────────────────────────────────────
        running = threading.Event()

        # ── log panel ──────────────────────────────────────────────────────────
        log_list = ft.ListView(
            expand=True,
            spacing=1,
            auto_scroll=True,
            padding=ft.Padding.symmetric(horizontal=10, vertical=6),
        )
        gui_logger = GuiLogger(log_list, page)

        # Stop the flush loop when the window is closed
        def _on_window_event(e):
            if e.data == "close":
                gui_logger.stop()
        page.on_window_event = _on_window_event

        def clear_log(_=None):
            log_list.controls.clear()
            page.update()

        _log_inner_container = ft.Container(
            content=log_list,
            bgcolor="#14FFFFFF",
            border_radius=8,
            expand=True,
        )
        _log_label = ft.Text("Log", size=13, weight=ft.FontWeight.BOLD, color="#AAAAAA")
        log_panel = ft.Container(
            content=ft.Column(
                [
                    ft.Row(
                        [
                            _log_label,
                            ft.IconButton(ft.Icons.DELETE_SWEEP_OUTLINED,
                                          icon_size=18, tooltip="Clear log",
                                          on_click=clear_log),
                        ],
                        alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                    ),
                    _log_inner_container,
                ],
                spacing=4,
                expand=True,
            ),
            padding=ft.Padding.symmetric(horizontal=12, vertical=8),
            expand=True,
        )

        # ── helpers ────────────────────────────────────────────────────────────
        def log(msg: str):
            gui_logger.write(msg + "\n")

        def set_running(state: bool):
            run_btn.disabled = state
            batch_run_btn.disabled = state
            progress_ring.visible = state
            sf_progress_bar.visible = state
            sf_progress_label.visible = state
            b_progress_bar.visible = state
            b_progress_label.visible = state
            if not state:
                sf_progress_bar.value = 0
                sf_progress_label.value = ""
                b_progress_bar.value = 0
                b_progress_label.value = ""
            page.update()

        def run_in_thread(fn, progress_bar=None, progress_label=None):
            if running.is_set():
                log("⚠ A process is already running.")
                return
            running.set()
            set_running(True)

            # ── Fixed milestones: (keyword_in_stdout, bar_value, label) ──────
            # These fire once when the matching text appears in a print() call.
            _steps = [
                ("Loading spaCy",                0.04, "Loading language model…"),
                ("Extracting nouns",             0.08, "Reading document…"),
                ("Analyzing text with spaCy",    0.12, "Splitting text into chunks…"),
                ("Pre-computing vectors",        0.30, "Pre-computing noun vectors…"),
                ("Building index using",         0.30, "Building hierarchical index…"),
                ("Building index by",            0.86, "Indexing document locations…"),
                ("Writing hierarchical index",   0.92, "Writing output file…"),
                ("Process completed",            1.00, "Done ✓"),
                ("Batch processing summary",     1.00, "Batch complete ✓"),
            ]

            # ── tqdm phase ranges: desc_substr → (bar_start, bar_end, label) ─
            # Each tqdm loop drives the bar smoothly from bar_start to bar_end
            # by reading the live percentage that tqdm embeds in every line.
            _tqdm_phases = {
                # extract_nouns: chunked nlp.pipe() — replaces old blocking nlp(text)
                "analyzing text":       (0.12, 0.28, "Analysing text with NLP…"),
                # _strategy_*: pre-vectorisation batch (the previously frozen step)
                "computing vectors":    (0.30, 0.72, "Computing semantic vectors…"),
                # _strategy_keywords / _strategy_frequent: similarity comparison
                "finding correlations": (0.72, 0.86, "Finding semantic correlations…"),
                # _strategy_hybrid: outer keyword loop
                "processing keywords":  (0.72, 0.86, "Building keyword index…"),
                # extract_occurrences: page / paragraph scan
                "indexing pages":       (0.86, 0.92, "Indexing pages…"),
                "indexing paragraphs":  (0.86, 0.92, "Indexing paragraphs…"),
                # _strategy_auto: clustering
                "creating clusters":    (0.30, 0.72, "Clustering nouns…"),
            }

            import re as _re_step

            def _on_step(line: str):
                if progress_bar is None:
                    return
                stripped = line.strip()
                l_lower  = stripped.lower()

                # Helper: push a progress update onto Flet's event loop so it
                # renders immediately, without waiting for the flush-loop batch.
                def _push_progress(bar_val: float, lbl_val: str):
                    _bar   = progress_bar
                    _lbl   = progress_label
                    _page  = page
                    async def _apply():
                        _bar.value = bar_val
                        if _lbl is not None:
                            _lbl.value = lbl_val
                        try:
                            await _page.update_async()
                        except Exception:
                            try:
                                _page.update()
                            except Exception:
                                pass
                    try:
                        page.run_task(_apply)
                    except Exception:
                        # Fallback for Flet builds without run_task
                        _bar.value = bar_val
                        if _lbl is not None:
                            _lbl.value = lbl_val
                        gui_logger.schedule_update()

                # ── tqdm lines: "  desc:  45%|████ | 45/100 [...]" ──────────
                # tqdm embeds the live percentage as "<N>%|" — parse it and
                # interpolate within the matching phase's bar range so the
                # progress bar moves continuously item-by-item.
                tqdm_match = _re_step.search(r'(\d+)%\|', stripped)
                if tqdm_match:
                    pct = int(tqdm_match.group(1)) / 100.0
                    for desc, (start, end, lbl) in _tqdm_phases.items():
                        if desc in l_lower:
                            _push_progress(round(start + pct * (end - start), 4), lbl)
                            return
                    return  # tqdm line from an unknown phase – ignore

                # ── Fixed milestones (non-tqdm print() calls) ────────────────
                for keyword, value, label in _steps:
                    if keyword.lower() in l_lower:
                        _push_progress(value, label)
                        break

            gui_logger.set_step_callback(_on_step)

            def _worker():
                old_out, old_err = sys.stdout, sys.stderr
                sys.stdout = sys.stderr = gui_logger
                try:
                    fn()
                except SystemExit:
                    pass
                except Exception as exc:
                    import traceback
                    log(f"✗ Unhandled error: {exc}")
                    log(traceback.format_exc())
                finally:
                    sys.stdout, sys.stderr = old_out, old_err
                    gui_logger.set_step_callback(None)
                    # Give the flush loop enough time to drain the queue AND
                    # let all run_task()-scheduled progress updates execute
                    # before we reset the controls.
                    _time.sleep(gui_logger.UI_REFRESH_INTERVAL * 4)
                    running.clear()
                    # Schedule the UI reset on Flet's event loop so it executes
                    # AFTER any still-pending run_task() progress updates,
                    # preventing a race where set_running(False) clears the bar
                    # before "Done ✓" is rendered.
                    async def _finish():
                        set_running(False)
                    try:
                        page.run_task(_finish)
                    except Exception:
                        set_running(False)

            threading.Thread(target=_worker, daemon=True).start()

        # ══════════════════════════════════════════════════════════════════════
        # SINGLE FILE TAB
        # ══════════════════════════════════════════════════════════════════════

        # ── shared styles ────────────────────────────────────────────────────
        _HINT_STYLE = ft.TextStyle(color="#777777", italic=True)

        # ── path / integer validation (on_change – avoids Flet InputFilter bugs) ─
        import re as _re, platform as _platform, os as _os

        _FORBIDDEN = _re.compile(r'[<>"|?*\x00-\x1f]')

        def _sanitize_path(field: ft.TextField) -> None:
            """Strip forbidden path characters as the user types."""
            raw = field.value or ""
            clean = _FORBIDDEN.sub("", raw)
            if clean != raw:
                field.value = clean
                field.update()

        def _validate_path(field: ft.TextField, must_exist: bool = False) -> None:
            """Show error_text on blur when the path is clearly invalid."""
            val = (field.value or "").strip()
            if not val:
                field.error_text = None
                field.update()
                return
            if _platform.system() == "Windows":
                if len(val) >= 2 and val[1] == ":" and not val[0].isalpha():
                    field.error_text = "Invalid drive letter"
                    field.update()
                    return
            if must_exist and not _os.path.exists(val):
                field.error_text = "File not found"
                field.update()
                return
            field.error_text = None
            field.update()

        def _sanitize_int(field: ft.TextField, fallback: str = "1") -> None:
            """Keep only digit characters; restore fallback if field is emptied."""
            raw = field.value or ""
            clean = "".join(c for c in raw if c.isdigit())
            if clean != raw:
                field.value = clean
                field.update()

        def _restore_int(field: ft.TextField, fallback: str = "1") -> None:
            """On blur, restore fallback value when field is empty."""
            if not (field.value or "").strip():
                field.value = fallback
                field.update()

        def _sanitize_min_occ(field: ft.TextField) -> None:
            """Keep only digit characters for min occurrences field."""
            raw = field.value or ""
            clean = "".join(c for c in raw if c.isdigit())
            if clean != raw:
                field.value = clean
                field.update()

        def _restore_min_occ(field: ft.TextField) -> None:
            """On blur, ensure min occurrences is at least 1."""
            val = (field.value or "").strip()
            if not val or not val.isdigit() or int(val) < 1:
                field.value = "1"
                field.update()


        # ── file pickers ──────────────────────────────────────────────────────────────────────────────────
        sf_input_field  = ft.TextField(
            label="Input file", expand=True, dense=True,
            hint_text="Select a .pdf, .docx, or .odt file",
            hint_style=_HINT_STYLE,
            on_change=lambda e: (_sanitize_path(sf_input_field), _sf_auto_mode()),
            on_blur=lambda e: _validate_path(sf_input_field, must_exist=True),
        )

        def _sf_auto_mode():
            val = (sf_input_field.value or "").strip().lower()
            if val.endswith((".docx", ".odt")):
                sf_mode.value = "paragraph"
                try:
                    sf_mode.update()
                except Exception:
                    pass
            # Disable "Mark document" when a PDF is selected
            is_pdf = val.endswith(".pdf")
            sf_mark.disabled = is_pdf
            if is_pdf:
                sf_mark.value = False
            try:
                sf_mark.update()
            except Exception:
                pass
        sf_output_field = ft.TextField(
            label="Output file", expand=True, dense=True,
            hint_text="e.g. index.txt",
            hint_style=_HINT_STYLE,
            on_change=lambda e: _sanitize_path(sf_output_field),
            on_blur=lambda e: _validate_path(sf_output_field, must_exist=False),
        )

        _sf_input_picker  = ft.FilePicker()
        _sf_output_picker = ft.FilePicker()

        async def sf_input_pick(_=None):
            files = await _sf_input_picker.pick_files(
                dialog_title="Select input document",
                allowed_extensions=["pdf", "docx", "odt"],
                allow_multiple=False,
            )
            if files:
                sf_input_field.value = files[0].path
                _sf_auto_mode()
                page.update()

        async def sf_output_pick(_=None):
            path = await _sf_output_picker.save_file(
                dialog_title="Save output index",
                file_name="index.txt",
                allowed_extensions=["txt", "json", "csv", "html"],
            )
            if path:
                sf_output_field.value = path
                page.update()

        # ── keywords ──────────────────────────────────────────────────────────
        sf_keywords = ft.TextField(
            label="Keywords (semicolon-separated)",
            hint_text="e.g. economy; law; justice",
            hint_style=_HINT_STYLE,
            expand=True, dense=True,
        )

        # ── strategy ──────────────────────────────────────────────────────────
        def on_strategy_change(e):
            kw_section.visible    = sf_strategy.value in ("keywords", "hybrid")
            auto_section.visible  = sf_strategy.value == "auto"
            hyb_section.visible   = sf_strategy.value == "hybrid"
            freq_section.visible  = sf_strategy.value == "frequent"
            page.update()

        sf_strategy = ft.Dropdown(
            label="Strategy",
            value="keywords",
            options=[
                ft.DropdownOption("keywords", "Keywords – 2 levels (default)"),
                ft.DropdownOption("auto",     "Auto clustering – 2 levels"),
                ft.DropdownOption("hybrid",   "Hybrid – 3 levels (keywords + sub-clusters)"),
                ft.DropdownOption("frequent", "Frequent terms – 2 levels"),
            ],
            dense=True,
            expand=True,
        )
        sf_strategy.on_change = on_strategy_change

        # ── threshold ─────────────────────────────────────────────────────────
        threshold_label = ft.Text("Similarity threshold: 0.50", size=12)

        def on_threshold_change(e):
            threshold_label.value = f"Similarity threshold: {sf_threshold.value:.2f}"
            page.update()

        sf_threshold = ft.Slider(min=0.0, max=1.0, value=0.5, divisions=20,
                                  on_change=on_threshold_change)

        # ── mode ──────────────────────────────────────────────────────────────
        sf_mode = ft.Dropdown(
            label="Indexing mode",
            value="page",
            options=[
                ft.DropdownOption("page",      "Page (PDF only)"),
                ft.DropdownOption("paragraph", "Paragraph"),
            ],
            dense=True,
            expand=True,
        )

        # ── export format ─────────────────────────────────────────────────────
        sf_format = ft.Dropdown(
            label="Export format",
            value="txt",
            options=[
                ft.DropdownOption("txt",  "Plain text (.txt)"),
                ft.DropdownOption("json", "JSON (.json)"),
                ft.DropdownOption("csv",  "CSV (.csv)"),
                ft.DropdownOption("html", "HTML (.html)"),
            ],
            dense=True,
            expand=True,
        )

        # ── spaCy model ───────────────────────────────────────────────────────
        def _get_installed_models() -> list:
            """Return only the spaCy models that are currently importable."""
            installed = []
            for m in LexiFinderCLI.SUPPORTED_MODELS:
                try:
                    __import__(m)
                    installed.append(m)
                except ImportError:
                    pass
            return installed

        _installed = _get_installed_models()
        _default_model = (
            LexiFinderCLI.DEFAULT_MODEL if LexiFinderCLI.DEFAULT_MODEL in _installed
            else (_installed[0] if _installed else None)
        )

        sf_model = ft.Dropdown(
            label="spaCy model",
            value=_default_model,
            options=(
                [ft.DropdownOption(m, f"{m}  –  {LexiFinderCLI.SUPPORTED_MODELS[m]}")
                 for m in _installed]
                if _installed
                else [ft.DropdownOption("", "No models installed – go to Models tab")]
            ),
            dense=True,
            expand=True,
        )
        # ── strategy-specific sections ────────────────────────────────────────
        sf_max_per_cat = ft.TextField(label="Max sub-entries / category",
                                       value="30", expand=True, dense=True,
                                       keyboard_type=ft.KeyboardType.NUMBER)

        kw_section = ft.Column([sf_keywords], visible=True)

        sf_clusters = ft.TextField(label="N. clusters", value="8",
                                    expand=True, dense=True,
                                    keyboard_type=ft.KeyboardType.NUMBER)
        auto_section = ft.Column(
            [ft.Row([sf_clusters], spacing=12)], visible=False
        )

        sf_subclusters = ft.TextField(label="N. sub-clusters", value="3",
                                       expand=True, dense=True,
                                       keyboard_type=ft.KeyboardType.NUMBER)
        hyb_section = ft.Column(
            [sf_keywords, ft.Row([sf_subclusters], spacing=12)], visible=False
        )

        sf_top = ft.TextField(label="Top N frequent terms", value="15",
                               expand=True, dense=True,
                               keyboard_type=ft.KeyboardType.NUMBER)
        freq_section = ft.Column(
            [ft.Row([sf_top], spacing=12)], visible=False
        )

        # ── filters ───────────────────────────────────────────────────────────
        sf_exclude_generic = ft.Checkbox(label="Exclude generic words", value=False)
        sf_min_occ = ft.TextField(label="Min occurrences", value="1",
                                   width=180, dense=True,
                                   keyboard_type=ft.KeyboardType.NUMBER,
                                   on_change=lambda e: _sanitize_min_occ(sf_min_occ),
                                   on_blur=lambda e: _restore_min_occ(sf_min_occ))
        sf_mark = ft.Checkbox(label="Mark document (DOCX/ODT only)", value=False)

        # ── config save/load ──────────────────────────────────────────────────
        def get_sf_config() -> dict:
            return {
                "keywords":        sf_keywords.value or "",
                "threshold":       sf_threshold.value,
                "mode":            sf_mode.value,
                "strategy":        sf_strategy.value,
                "max_per_category": int(sf_max_per_cat.value or 30),
                "clusters":        int(sf_clusters.value or 8),
                "subclusters":     int(sf_subclusters.value or 3),
                "top":             int(sf_top.value or 15),
                "exclude_generic": sf_exclude_generic.value,
                "min_occurrences": int(sf_min_occ.value or 1),
                "model":           sf_model.value,
                "export_format":   sf_format.value,
                "mark":            sf_mark.value,
            }

        def apply_sf_config(cfg: dict):
            if "keywords"        in cfg: sf_keywords.value        = cfg["keywords"]
            if "threshold"       in cfg:
                sf_threshold.value = float(cfg["threshold"])
                threshold_label.value = f"Similarity threshold: {sf_threshold.value:.2f}"
            if "mode"            in cfg: sf_mode.value            = cfg["mode"]
            if "strategy"        in cfg:
                sf_strategy.value = cfg["strategy"]
                on_strategy_change(None)
            if "max_per_category" in cfg: sf_max_per_cat.value   = str(cfg["max_per_category"])
            if "clusters"        in cfg: sf_clusters.value        = str(cfg["clusters"])
            if "subclusters"     in cfg: sf_subclusters.value     = str(cfg["subclusters"])
            if "top"             in cfg: sf_top.value             = str(cfg["top"])
            if "exclude_generic" in cfg: sf_exclude_generic.value = bool(cfg["exclude_generic"])
            if "min_occurrences" in cfg: sf_min_occ.value         = str(cfg["min_occurrences"])
            if "model"           in cfg: sf_model.value           = cfg["model"]
            if "export_format"   in cfg: sf_format.value          = cfg["export_format"]
            if "mark"            in cfg: sf_mark.value            = bool(cfg["mark"])
            page.update()

        _save_config_picker = ft.FilePicker()
        _load_config_picker = ft.FilePicker()

        async def save_config_dialog(_=None):
            path = await _save_config_picker.save_file(
                dialog_title="Save configuration",
                file_name="lexifinder-config.json",
                allowed_extensions=["json"],
            )
            if path:
                _save_config(path)


        def _save_config(path: str):
            cfg = get_sf_config()
            try:
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(cfg, f, indent=2, ensure_ascii=False)
                log(f"✓ Configuration saved to '{path}'")
            except Exception as exc:
                log(f"✗ Error saving configuration: {exc}")

        async def load_config_dialog(_=None):
            files = await _load_config_picker.pick_files(
                dialog_title="Load configuration",
                allowed_extensions=["json"],
                allow_multiple=False,
            )
            if files:
                _load_config(files[0].path)


        def _load_config(path: str):
            try:
                with open(path, "r", encoding="utf-8") as f:
                    cfg = json.load(f)
                apply_sf_config(cfg)
                log(f"✓ Configuration loaded from '{path}'")
            except Exception as exc:
                log(f"✗ Error loading configuration: {exc}")

        # ── run single file ───────────────────────────────────────────────────
        def run_single(_=None):
            inp  = sf_input_field.value.strip()
            outp = sf_output_field.value.strip()
            if not inp:
                log("✗ Please specify an input file."); return
            if not outp:
                log("✗ Please specify an output file."); return

            strategy  = sf_strategy.value
            keywords  = None
            if strategy in ("keywords", "hybrid"):
                raw_kw = sf_keywords.value.strip()
                if not raw_kw:
                    log(f"✗ Strategy '{strategy}' requires keywords."); return
                try:
                    keywords = parse_keywords(raw_kw)
                except ValueError as e:
                    log(f"✗ {e}"); return

            # fix extension
            fmt = sf_format.value
            base, _ = os.path.splitext(outp)
            outp = f"{base}.{fmt}"

            opts = dict(
                max_per_category = int(sf_max_per_cat.value or 30),
                clusters         = int(sf_clusters.value  or 8),
                subclusters      = int(sf_subclusters.value or 3),
                top              = int(sf_top.value or 15),
                exclude_generic  = sf_exclude_generic.value,
                min_occurrences  = int(sf_min_occ.value or 1),
            )
            model     = sf_model.value
            threshold = float(sf_threshold.value)
            mode      = sf_mode.value
            mark      = sf_mark.value

            def _run():
                log("▶ Single File Processing started")
                log(f"  Input:    {inp}")
                log(f"  Output:   {outp}")
                log(f"  Strategy: {strategy}  |  Model: {model}")
                finder = LexiFinderCLI(similarity_threshold=threshold, model_name=model)
                finder.process(
                    inp, outp,
                    keywords=keywords,
                    mode=mode,
                    mark_document=mark,
                    strategy=strategy,
                    export_format=fmt,
                    **opts,
                )

            run_in_thread(_run, progress_bar=sf_progress_bar, progress_label=sf_progress_label)

        progress_ring = ft.ProgressRing(width=20, height=20, stroke_width=2, visible=False)
        sf_progress_bar   = ft.ProgressBar(value=0, visible=False, color="#2196F3",
                                            bgcolor="#333333", expand=True)
        sf_progress_label = ft.Text("", size=11, color="#90CAF9", visible=False)
        run_btn = ft.FilledButton("▶  Run", on_click=run_single, bgcolor="#2196F3")

        single_tab_content = ft.Column(
            [
                # Row 1: files
                ft.Row([
                    sf_input_field,
                    ft.IconButton(ft.Icons.FOLDER_OPEN, tooltip="Browse input",
                                  on_click=sf_input_pick),
                ], spacing=4),
                ft.Row([
                    sf_output_field,
                    ft.IconButton(ft.Icons.SAVE_AS, tooltip="Browse output",
                                  on_click=sf_output_pick),
                ], spacing=4),
                ft.Divider(height=1),
                # Row 2: strategy + model
                ft.Row([sf_strategy, sf_model], spacing=12),
                kw_section, auto_section, hyb_section, freq_section,
                ft.Divider(height=1),
                # Row 3: threshold
                ft.Column([threshold_label, sf_threshold], spacing=0),
                ft.Divider(height=1),
                # Row 4: mode + format + max
                ft.Row([sf_mode, sf_format, sf_max_per_cat], spacing=12),
                ft.Divider(height=1),
                # Row 5: filters
                ft.Row([sf_exclude_generic, sf_min_occ, sf_mark], spacing=20,
                       wrap=True),
                ft.Divider(height=1),
                # Row 6: config + run
                ft.Row([
                    ft.OutlinedButton("💾 Save config", on_click=save_config_dialog),
                    ft.OutlinedButton("📂 Load config", on_click=load_config_dialog),
                    ft.Container(expand=True),
                    progress_ring,
                    run_btn,
                ], alignment=ft.MainAxisAlignment.END, spacing=8),
                # Progress bar (visible while running)
                ft.Row([sf_progress_bar], visible=True),
                sf_progress_label,
            ],
            spacing=8,
            scroll=ft.ScrollMode.AUTO,
            expand=True,
        )

        # ══════════════════════════════════════════════════════════════════════
        # BATCH MODE TAB
        # ══════════════════════════════════════════════════════════════════════

        b_input_dir  = ft.TextField(
            label="Input directory", expand=True, dense=True,
            on_change=lambda e: _sanitize_path(b_input_dir),
            on_blur=lambda e: _validate_path(b_input_dir, must_exist=True),
        )
        b_output_dir = ft.TextField(
            label="Output directory", expand=True, dense=True,
            on_change=lambda e: _sanitize_path(b_output_dir),
            on_blur=lambda e: _validate_path(b_output_dir, must_exist=False),
        )

        _b_input_picker  = ft.FilePicker()
        _b_output_picker = ft.FilePicker()

        async def b_input_pick(_=None):
            path = await _b_input_picker.get_directory_path(dialog_title="Select input directory")
            if path:
                b_input_dir.value = path
                page.update()

        async def b_output_pick(_=None):
            path = await _b_output_picker.get_directory_path(dialog_title="Select output directory")
            if path:
                b_output_dir.value = path
                page.update()

        def _b_update_mark(_=None):
            pat = (b_pattern.value or "").strip().lower()
            # Disable mark when the pattern exclusively targets PDF files
            is_pdf_only = pat.endswith(".pdf")
            b_mark.disabled = is_pdf_only
            if is_pdf_only:
                b_mark.value = False
            try:
                b_mark.update()
            except Exception:
                pass

        b_pattern  = ft.TextField(label="File pattern", value="*",
                                   hint_text="e.g. *.pdf or chapter*.docx",
                                   hint_style=_HINT_STYLE,
                                   width=200, dense=True,
                                   on_change=_b_update_mark)
        b_keywords = ft.TextField(label="Keywords (semicolon-separated)",
                                   hint_text="e.g. economy; law; justice",
                                   hint_style=_HINT_STYLE,
                                   expand=True, dense=True)

        def on_b_strategy_change(e):
            b_kw_section.visible   = b_strategy.value in ("keywords", "hybrid")
            b_auto_section.visible = b_strategy.value == "auto"
            b_hyb_section.visible  = b_strategy.value == "hybrid"
            b_freq_section.visible = b_strategy.value == "frequent"
            page.update()

        b_strategy = ft.Dropdown(
            label="Strategy", value="keywords",
            options=[
                ft.DropdownOption("keywords", "Keywords – 2 levels (default)"),
                ft.DropdownOption("auto",     "Auto clustering – 2 levels"),
                ft.DropdownOption("hybrid",   "Hybrid – 3 levels (keywords + sub-clusters)"),
                ft.DropdownOption("frequent", "Frequent terms – 2 levels"),
            ],
            dense=True,
            expand=True,
        )
        b_strategy.on_change = on_b_strategy_change
        b_model = ft.Dropdown(
            label="spaCy model",
            value=_default_model,
            options=(
                [ft.DropdownOption(m, f"{m}  –  {LexiFinderCLI.SUPPORTED_MODELS[m]}")
                 for m in _installed]
                if _installed
                else [ft.DropdownOption("", "No models installed – go to Models tab")]
            ),
            dense=True,
            expand=True,
        )
        b_threshold_label = ft.Text("Similarity threshold: 0.50", size=12)

        def on_b_threshold_change(e):
            b_threshold_label.value = f"Similarity threshold: {b_threshold.value:.2f}"
            page.update()

        b_threshold   = ft.Slider(min=0.0, max=1.0, value=0.5, divisions=20,
                                   on_change=on_b_threshold_change)
        b_mode        = ft.Dropdown(label="Indexing mode", value="page",
                                     options=[ft.DropdownOption("page",      "Page (PDF only)"),
                                              ft.DropdownOption("paragraph", "Paragraph")],
                                     dense=True, expand=True)
        b_format      = ft.Dropdown(label="Export format", value="txt",
                                     options=[ft.DropdownOption("txt",  "Plain text (.txt)"),
                                              ft.DropdownOption("json", "JSON (.json)"),
                                              ft.DropdownOption("csv",  "CSV (.csv)"),
                                              ft.DropdownOption("html", "HTML (.html)")],
                                     dense=True, expand=True)
        b_max_per_cat = ft.TextField(label="Max sub-entries / category", value="30",
                                      expand=True, dense=True,
                                      keyboard_type=ft.KeyboardType.NUMBER)
        b_clusters    = ft.TextField(label="N. clusters", value="8", expand=True, dense=True,
                                      keyboard_type=ft.KeyboardType.NUMBER)
        b_subclusters = ft.TextField(label="N. sub-clusters", value="3", expand=True, dense=True,
                                      keyboard_type=ft.KeyboardType.NUMBER)
        b_top         = ft.TextField(label="Top N frequent terms", value="15", expand=True, dense=True,
                                      keyboard_type=ft.KeyboardType.NUMBER)
        b_exclude_generic = ft.Checkbox(label="Exclude generic words", value=False)
        b_min_occ     = ft.TextField(label="Min occurrences", value="1",
                                      width=180, dense=True,
                                      keyboard_type=ft.KeyboardType.NUMBER,
                                      on_change=lambda e: _sanitize_min_occ(b_min_occ),
                                      on_blur=lambda e: _restore_min_occ(b_min_occ))
        b_mark        = ft.Checkbox(label="Mark documents (DOCX/ODT)", value=False)

        b_mode_note = ft.Text(
            "ℹ  Mixed batch: Page mode applies to PDFs; DOCX/ODT always use Paragraph mode.",
            size=11, color="#90CAF9", visible=False,
        )

        def _on_b_mode_change(e):
            b_mode_note.visible = (b_mode.value == "page")
            b_mode_note.update()

        b_mode.on_change = _on_b_mode_change

        b_kw_section   = ft.Column([b_keywords], visible=True)
        b_auto_section = ft.Column(
            [ft.Row([b_clusters], spacing=12)], visible=False
        )
        b_hyb_section  = ft.Column(
            [b_keywords, ft.Row([b_subclusters], spacing=12)], visible=False
        )
        b_freq_section = ft.Column(
            [ft.Row([b_top], spacing=12)], visible=False
        )

        def run_batch(_=None):
            inp  = b_input_dir.value.strip()
            outp = b_output_dir.value.strip()
            if not inp:
                log("✗ Please specify an input directory."); return
            if not outp:
                log("✗ Please specify an output directory."); return

            strategy = b_strategy.value
            keywords = None
            if strategy in ("keywords", "hybrid"):
                raw_kw = b_keywords.value.strip()
                if not raw_kw:
                    log(f"✗ Strategy '{strategy}' requires keywords."); return
                try:
                    keywords = parse_keywords(raw_kw)
                except ValueError as e:
                    log(f"✗ {e}"); return

            opts = dict(
                max_per_category = int(b_max_per_cat.value or 30),
                clusters         = int(b_clusters.value or 8),
                subclusters      = int(b_subclusters.value or 3),
                top              = int(b_top.value or 15),
                exclude_generic  = b_exclude_generic.value,
                min_occurrences  = int(b_min_occ.value or 1),
            )
            model   = b_model.value
            thresh  = float(b_threshold.value)
            mode    = b_mode.value
            fmt     = b_format.value
            pattern = b_pattern.value.strip() or "*"
            mark    = b_mark.value

            def _run():
                log("▶ Batch Processing started")
                finder = LexiFinderCLI(similarity_threshold=thresh, model_name=model)
                finder.process_batch(
                    inp, outp,
                    pattern=pattern,
                    keywords=keywords,
                    mode=mode,
                    mark_document=mark,
                    strategy=strategy,
                    export_format=fmt,
                    **opts,
                )

            run_in_thread(_run, progress_bar=b_progress_bar, progress_label=b_progress_label)

        b_progress_bar   = ft.ProgressBar(value=0, visible=False, color="#2196F3",
                                           bgcolor="#333333", expand=True)
        b_progress_label = ft.Text("", size=11, color="#90CAF9", visible=False)
        batch_run_btn = ft.FilledButton("▶  Run Batch", on_click=run_batch, bgcolor="#2196F3")

        batch_tab_content = ft.Column(
            [
                ft.Row([b_input_dir,
                        ft.IconButton(ft.Icons.FOLDER_OPEN, tooltip="Browse",
                                      on_click=b_input_pick)], spacing=4),
                ft.Row([b_output_dir,
                        ft.IconButton(ft.Icons.FOLDER_OPEN, tooltip="Browse",
                                      on_click=b_output_pick)], spacing=4),
                ft.Row([b_pattern, b_strategy, b_model], spacing=12),
                b_kw_section, b_auto_section, b_hyb_section, b_freq_section,
                ft.Divider(height=1),
                ft.Column([b_threshold_label, b_threshold], spacing=0),
                ft.Divider(height=1),
                ft.Row([b_mode, b_format, b_max_per_cat], spacing=12),
                b_mode_note,
                ft.Divider(height=1),
                ft.Row([b_exclude_generic, b_min_occ, b_mark], spacing=20, wrap=True),
                ft.Divider(height=1),
                ft.Row([
                    ft.Container(expand=True),
                    batch_run_btn,
                ], alignment=ft.MainAxisAlignment.END),
                # Progress bar (visible while running)
                ft.Row([b_progress_bar], visible=True),
                b_progress_label,
            ],
            spacing=8,
            scroll=ft.ScrollMode.AUTO,
            expand=True,
        )

        # ══════════════════════════════════════════════════════════════════════
        # MODELS TAB
        # ══════════════════════════════════════════════════════════════════════

        def _check_installed():
            result = []
            for m in LexiFinderCLI.SUPPORTED_MODELS:
                try:
                    __import__(m)
                    result.append(m)
                except ImportError:
                    pass
            return result

        download_list_col   = ft.Column(spacing=4, scroll=ft.ScrollMode.AUTO)
        installed_list_col  = ft.Column(spacing=4, scroll=ft.ScrollMode.AUTO)

        # Mutable holder so refresh_model_lists always picks the current palette
        _current_pal = [None]   # will be set by _apply_theme on first call
        _page_ready   = [False] # True only after page.add() has been called

        def refresh_model_lists():
            inst = _check_installed()
            not_inst = [m for m in LexiFinderCLI.SUPPORTED_MODELS if m not in inst]

            download_list_col.controls.clear()
            if not_inst:
                for m in not_inst:
                    desc = LexiFinderCLI.SUPPORTED_MODELS[m]
                    mn = m  # capture
                    def make_dl_btn(model_name):
                        def do_download(_=None):
                            def _run():
                                LexiFinderCLI.download_model(model_name)
                                refresh_model_lists()
                                page.update()
                            run_in_thread(_run)
                        return ft.FilledButton("⬇  Download", on_click=do_download,
                                               style=ft.ButtonStyle(bgcolor="#1565C0"))
                    download_list_col.controls.append(
                        ft.Row([
                            ft.Text(f"{m}", size=12, expand=2, selectable=True),
                            ft.Text(desc, size=11, color="#AAAAAA", expand=3),
                            make_dl_btn(mn),
                        ], spacing=8)
                    )
            else:
                download_list_col.controls.append(
                    ft.Text("All supported models are already installed.",
                            size=12, color="#AAAAAA", italic=True)
                )

            installed_list_col.controls.clear()
            if inst:
                for m in inst:
                    desc = LexiFinderCLI.SUPPORTED_MODELS.get(m, "")
                    bundled = (m == LexiFinderCLI.DEFAULT_MODEL
                               and getattr(sys, "frozen", False))
                    mn = m
                    def make_del_btn(model_name):
                        def do_delete(_=None):
                            def confirm_delete(e):
                                dlg.open = False
                                page.update()
                                if e.control.text == "Yes, delete":
                                    def _run():
                                        LexiFinderCLI.delete_model(model_name, force=True)
                                        refresh_model_lists()
                                        page.update()
                                    run_in_thread(_run)
                            dlg = ft.AlertDialog(
                                title=ft.Text("Confirm deletion"),
                                content=ft.Text(f"Delete model '{model_name}'?"),
                                actions=[
                                    ft.TextButton("Cancel",      on_click=confirm_delete),
                                    ft.TextButton("Yes, delete", on_click=confirm_delete),
                                ],
                            )
                            page.open(dlg)
                        red = (_current_pal[0] or {}).get("support_red", "#CF6679")
                        return ft.FilledButton("🗑  Delete", on_click=do_delete,
                                               style=ft.ButtonStyle(bgcolor=red))
                    tag = ft.Text("[bundled]", size=10, color="#90CAF9") if bundled else ft.Text("")
                    installed_list_col.controls.append(
                        ft.Row([
                            ft.Text(f"✓  {m}", size=12, expand=2, selectable=True),
                            ft.Text(desc, size=11, color="#AAAAAA", expand=3),
                            tag,
                            make_del_btn(mn),
                        ], spacing=8)
                    )
            else:
                installed_list_col.controls.append(
                    ft.Text("No models installed.", size=12, color="#AAAAAA", italic=True)
                )

            if _page_ready[0]:
                download_list_col.update()
                installed_list_col.update()

        # populate on first render – fired after page.add() sets _page_ready
        import threading as _threading

        _models_desc_text = ft.Text(
            "The default bundled model is en_core_web_md (English, medium). "
            "Additional models can be downloaded for other languages.",
            size=12, color="#AAAAAA"
        )
        models_tab_content = ft.Column(
            [
                ft.Text("Model Management", size=14, weight=ft.FontWeight.BOLD),
                _models_desc_text,
                ft.Divider(height=1),
                ft.Text("Models available for download", size=13, weight=ft.FontWeight.W_600),
                download_list_col,
                ft.Divider(height=1),
                ft.Text("Installed models", size=13, weight=ft.FontWeight.W_600),
                installed_list_col,
            ],
            spacing=10,
            expand=True,
            scroll=ft.ScrollMode.AUTO,
        )

        # ══════════════════════════════════════════════════════════════════════
        # ══════════════════════════════════════════════════════════════════════
        # ABOUT TAB
        # ══════════════════════════════════════════════════════════════════════

        def _open_url(url):
            import webbrowser
            webbrowser.open(url)

        _support_icon  = ft.Icon(ft.Icons.FAVORITE, size=20, color="#EF5350")
        _support_title = ft.Text("Support LexiFinder", size=13,
                                  weight=ft.FontWeight.W_600, color="#EF5350")
        _support_body  = ft.Text(
            "If you find LexiFinder useful, consider supporting its development "
            "with a small donation. Every contribution is greatly appreciated!",
            size=12, color="#AAAAAA",
        )
        _about_support_container = ft.Container(
            content=ft.Column([
                ft.Row([_support_icon, _support_title], spacing=8),
                _support_body,
                ft.Container(height=4),
                ft.FilledButton(
                    "💛  Donate via PayPal",
                    on_click=lambda _: _open_url(APP_PAYPAL),
                    style=ft.ButtonStyle(bgcolor="#1565C0"),
                ),
            ], spacing=8),
            padding=ft.Padding.all(16),
            border_radius=8,
            bgcolor="#1A1A2E",
        )
        _about_icon       = ft.Image(src="lexifinder.svg", width=52, height=52)
        _about_title      = ft.Text("LexiFinder", size=28, weight=ft.FontWeight.BOLD,
                                    color="#90CAF9")
        _about_version    = ft.Text(f"Version {APP_VERSION}", size=14, color="#AAAAAA")
        _about_tagline    = ft.Text("Hierarchical Analytic Index Generator",
                                    size=14, color="#CCCCCC")
        _about_author     = ft.Text(f"by {APP_AUTHOR}", size=13, color="#AAAAAA")
        _about_code_icon  = ft.Icon(ft.Icons.CODE, size=16, color="#90CAF9")
        _about_github_btn = ft.TextButton(
            "GitHub Repository",
            on_click=lambda _: _open_url(APP_REPO),
            style=ft.ButtonStyle(color="#90CAF9"),
        )
        _about_lang_icon  = ft.Icon(ft.Icons.LANGUAGE, size=16, color="#90CAF9")
        _about_web_btn    = ft.TextButton(
            "Author Webpage",
            on_click=lambda _: _open_url(APP_WEB),
            style=ft.ButtonStyle(color="#90CAF9"),
        )
        about_tab_content = ft.Column(
            [
                ft.Container(height=16),
                ft.Row([
                    _about_icon,
                    ft.Column([_about_title, _about_version], spacing=2),
                ], spacing=16, vertical_alignment=ft.CrossAxisAlignment.CENTER),
                _about_tagline,
                _about_author,
                ft.Divider(height=16),
                ft.Text("Links", size=13, weight=ft.FontWeight.W_600),
                ft.Row([_about_code_icon, _about_github_btn], spacing=4),
                ft.Row([_about_lang_icon, _about_web_btn], spacing=4),
                ft.Divider(height=16),
                _about_support_container,
            ],
            spacing=10,
            expand=True,
            scroll=ft.ScrollMode.AUTO,
        )

        # ══════════════════════════════════════════════════════════════════════
        # LAYOUT  –  NavigationBar replaces broken ft.Tabs
        # ══════════════════════════════════════════════════════════════════════

        _tab_bodies = [
            ft.Container(content=single_tab_content, padding=ft.Padding.all(16), expand=True, visible=True),
            ft.Container(content=batch_tab_content,  padding=ft.Padding.all(16), expand=True, visible=False),
            ft.Container(content=models_tab_content, padding=ft.Padding.all(16), expand=True, visible=False),
            ft.Container(content=about_tab_content,  padding=ft.Padding.all(16), expand=True, visible=False),
        ]
        _tab_body_stack = ft.Stack(controls=_tab_bodies, expand=True)

        def _on_nav_change(e):
            idx = nav_bar.selected_index
            for i, body in enumerate(_tab_bodies):
                body.visible = (i == idx)
            _tab_body_stack.update()

        nav_bar = ft.NavigationBar(
            selected_index=0,
            on_change=_on_nav_change,
            bgcolor="#1E1E1E",
            indicator_color="#1565C0",
            destinations=[
                ft.NavigationBarDestination(
                    icon=ft.Icons.ARTICLE_OUTLINED,
                    selected_icon=ft.Icons.ARTICLE,
                    label="Single File",
                ),
                ft.NavigationBarDestination(
                    icon=ft.Icons.FOLDER_COPY_OUTLINED,
                    selected_icon=ft.Icons.FOLDER_COPY,
                    label="Batch Mode",
                ),
                ft.NavigationBarDestination(
                    icon=ft.Icons.EXTENSION_OUTLINED,
                    selected_icon=ft.Icons.EXTENSION,
                    label="Models",
                ),
                ft.NavigationBarDestination(
                    icon=ft.Icons.INFO_OUTLINED,
                    selected_icon=ft.Icons.INFO,
                    label="About",
                ),
            ],
        )

        _header_title    = ft.Text("LexiFinder", size=20, weight=ft.FontWeight.BOLD,
                                  color="#90CAF9")
        _header_subtitle = ft.Text("Hierarchical Analytic Index Generator", size=13,
                                   color="#AAAAAA")
        _header_version  = ft.Text(f"v{APP_VERSION}", size=11, color="#AAAAAA")
        header = ft.Container(
            content=ft.Row([
                _header_title,
                _header_subtitle,
                ft.Container(expand=True),
                _header_version,
            ], spacing=10),
            padding=ft.Padding.symmetric(horizontal=16, vertical=10),
            bgcolor="#0F90CAF9",
        )

        page.add(
            ft.Column(
                [
                    header,
                    ft.Row(
                        [
                            # left: NavigationBar + content
                            ft.Container(
                                content=ft.Column([
                                    nav_bar,
                                    _tab_body_stack,
                                ], spacing=0, expand=True),
                                expand=3,
                            ),
                            ft.VerticalDivider(width=1),
                            # right: log
                            ft.Container(content=log_panel, expand=2),
                        ],
                        expand=True,
                        spacing=0,
                    ),
                ],
                spacing=0,
                expand=True,
            )
        )

        # Controls are now part of the page – safe to call .update() on them
        _page_ready[0] = True

        # ══════════════════════════════════════════════════════════════════════
        # ADAPTIVE THEME  –  dark / light mode following system preference
        # ══════════════════════════════════════════════════════════════════════

        _DARK = {
            "page_bg":        "#121212",
            "nav_bg":         "#1E1E1E",
            "log_inner_bg":   "#14FFFFFF",
            "header_bg":      "#1A237E22",
            "support_bg":     "#1A1A2E",
            "progress_bg":    "#333333",
            "muted":          "#AAAAAA",
            "accent":         "#90CAF9",
            "tagline":        "#CCCCCC",
            "support_red":    "#EF5350",
        }
        _LIGHT = {
            "page_bg":        "#F5F5F5",
            "nav_bg":         "#E8EAF6",
            "log_inner_bg":   "#FFFFFF",
            "header_bg":      "#BBDEFB",
            "support_bg":     "#E3F2FD",
            "progress_bg":    "#BDBDBD",
            "muted":          "#444444",
            "accent":         "#1565C0",
            "tagline":        "#333333",
            "support_red":    "#C62828",
        }

        def _apply_theme(_=None):
            is_dark = (page.platform_brightness == ft.Brightness.DARK
                       or page.theme_mode == ft.ThemeMode.DARK)
            pal = _DARK if is_dark else _LIGHT
            _current_pal[0] = pal
            refresh_model_lists()

            # Containers / backgrounds
            page.bgcolor                        = pal["page_bg"]
            nav_bar.bgcolor                     = pal["nav_bg"]
            _log_inner_container.bgcolor        = pal["log_inner_bg"]
            header.bgcolor                      = pal["header_bg"]
            _about_support_container.bgcolor    = pal["support_bg"]
            sf_progress_bar.bgcolor             = pal["progress_bg"]
            b_progress_bar.bgcolor              = pal["progress_bg"]

            # Header texts
            _header_title.color                 = pal["accent"]
            _header_subtitle.color              = pal["muted"]
            _header_version.color               = pal["muted"]

            # Log label
            _log_label.color                    = pal["muted"]

            # Progress / note labels
            sf_progress_label.color             = pal["accent"]
            b_progress_label.color              = pal["accent"]
            b_mode_note.color                   = pal["accent"]

            # About tab
            # _about_icon is an ft.Image – no colour to update
            _about_title.color                  = pal["accent"]
            _about_version.color                = pal["muted"]
            _about_tagline.color                = pal["tagline"]
            _about_author.color                 = pal["muted"]
            _about_code_icon.color              = pal["accent"]
            _about_lang_icon.color              = pal["accent"]
            _about_github_btn.style             = ft.ButtonStyle(color=pal["accent"])
            _about_web_btn.style                = ft.ButtonStyle(color=pal["accent"])

            # Models tab
            _models_desc_text.color             = pal["muted"]

            # About – support box
            _support_icon.color                 = pal["support_red"]
            _support_title.color                = pal["support_red"]
            _support_body.color                 = pal["muted"]

            # Hint style (update globally via page theme)
            page.theme = ft.Theme(color_scheme_seed=pal["accent"])

            page.update()

        page.on_platform_brightness_change = _apply_theme
        _apply_theme()  # apply on startup

        # First population of model lists (page is ready, .update() is safe)
        refresh_model_lists()

        log("LexiFinder GUI ready. Configure your options and click ▶ Run.")

    ft.run(main_gui)


if __name__ == "__main__":
    # Launch GUI if --gui flag is present or no arguments given (double-click)
    if "--gui" in sys.argv or (len(sys.argv) == 1 and not sys.stdin.isatty()
                                or len(sys.argv) == 1):
        run_gui()
    else:
        main()
